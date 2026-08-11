"""Text extraction from uploaded files.

Supported: PDF (pypdf), DOCX (python-docx), ODT (odfpy), TXT (UTF-8 with
fallback), Markdown (raw source kept verbatim so nothing is lost for the
viewer).
"""

from io import BytesIO

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from pypdf import PdfReader


def extract_text(content: bytes, file_type: str) -> str:
    if file_type == "pdf":
        return _extract_pdf(content)
    if file_type == "docx":
        return _extract_docx(content)
    if file_type == "odt":
        return _extract_odt(content)
    if file_type in ("txt", "md"):
        return _extract_text_file(content)
    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail=f"Unsupported file type '{file_type}'",
    )


def _extract_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages).strip()
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to extract text from PDF (corrupted or password-protected?)",
        )


def _extract_docx(content: bytes) -> str:
    try:
        doc = DocxDocument(BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to extract text from DOCX (corrupted file?)",
        )


def _extract_text_file(content: bytes) -> str:
    """Plain-text files (TXT / Markdown): decode, keep the content intact."""
    for encoding in ("utf-8", "windows-1251", "cp1252"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace").strip()


# ODF elements that terminate a text block (paragraph, heading, list item,
# table cell). Everything inside is extracted via teletype.extractText, which
# resolves spaces, tabs, line breaks, spans and hyperlinks.
_ODT_BLOCK_TAGS = frozenset(
    {
        "text:p",
        "text:h",
        "text:list-item",
        "table:table-cell",
        "table:covered-table-cell",
    }
)


def _extract_odt(content: bytes) -> str:
    """Extract the plain text of an ODT document (OpenDocument Text).

    Uses the lightweight, pure-Python ``odfpy`` package: the document is
    loaded and walked in document order, one line per paragraph / heading /
    list item / table cell. The result feeds the existing chunking → embedding
    pipeline unchanged.
    """
    from odf import teletype
    from odf.opendocument import load

    try:
        doc = load(BytesIO(content))
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="Failed to read ODT (corrupted file?)",
        )

    lines: list[str] = []

    def walk(node) -> None:
        tag = getattr(node, "tagName", "")
        if tag in _ODT_BLOCK_TAGS:
            text = " ".join(teletype.extractText(node).split())
            if text:
                lines.append(text)
            return  # block boundaries stop the recursion (nested blocks live here)
        for child in getattr(node, "childNodes", ()):
            walk(child)

    walk(doc.body)
    return "\n".join(lines)


def sniff_file_type(content: bytes) -> str | None:
    """Best-effort MIME detection from magic bytes. Returns None if unknown."""
    if content.startswith(b"%PDF-"):
        return "pdf"
    if content.startswith(b"PK\x03\x04"):
        return _sniff_zip(content)
    return None


def _sniff_zip(content: bytes) -> str | None:
    """DOCX and ODT are both ZIP containers; distinguish them by their entries.

    ODF documents carry a stored ``mimetype`` entry (``application/vnd.oasis
    .opendocument.text``), which Office Open XML files do not have.
    """
    from zipfile import BadZipFile, ZipFile

    try:
        with ZipFile(BytesIO(content)) as zf:
            names = zf.namelist()
            if "mimetype" in names:
                mime = zf.read("mimetype").decode("utf-8", "replace")
                if "opendocument" in mime or mime.startswith("application/vnd.oasis"):
                    return "odt"
            return "docx"
    except BadZipFile:
        return "docx"
