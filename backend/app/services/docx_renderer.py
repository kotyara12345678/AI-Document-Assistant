"""Render a validated DocumentSpec into a real .docx file.

python-docx writes OOXML as UTF-8 XML, so Cyrillic is preserved as-is and
empty / very long texts do not break the writer. Headings, paragraphs, bullet
and numbered lists and tables are emitted from the shared DocumentSpec, so
this module contains only DOCX-specific formatting — never content logic.
"""

from io import BytesIO

from docx import Document as DocxDocument

from app.schemas.document_spec import (
    DocumentSpec,
    HeadingBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
)


def render_docx(spec: DocumentSpec) -> bytes:
    """Return the bytes of a .docx built from ``spec``."""
    doc = DocxDocument()

    properties = doc.core_properties
    properties.title = spec.title
    if spec.metadata.author:
        properties.author = spec.metadata.author
    if spec.metadata.subject:
        properties.subject = spec.metadata.subject
    if spec.metadata.keywords:
        properties.keywords = spec.metadata.keywords

    doc.add_heading(spec.title, level=0)

    for block in spec.blocks:
        _render_block(doc, block)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_block(doc: DocxDocument, block) -> None:
    if isinstance(block, HeadingBlock):
        doc.add_heading(block.text, level=block.level)
    elif isinstance(block, ParagraphBlock):
        doc.add_paragraph(block.text)
    elif isinstance(block, ListBlock):
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            doc.add_paragraph(item, style=style)
    elif isinstance(block, TableBlock):
        _render_table(doc, block)


def _render_table(doc: DocxDocument, block: TableBlock) -> None:
    column_count = len(block.headers)
    if column_count == 0:
        column_count = max((len(row) for row in block.rows), default=0)
    if column_count == 0:
        return

    table = doc.add_table(rows=1 + len(block.rows), cols=column_count)
    table.style = "Table Grid"
    for index, header in enumerate(block.headers):
        table.cell(0, index).text = header
    for row_index, row in enumerate(block.rows, start=1):
        for cell_index in range(column_count):
            value = row[cell_index] if cell_index < len(row) else ""
            table.cell(row_index, cell_index).text = value
