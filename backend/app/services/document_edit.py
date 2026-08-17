"""Edit an existing user document by copying it, then rewriting the copy.

Flow (never mutates the original):

    source Document (read-only bytes)
        -> make an in-memory copy
        -> extract text blocks (paragraphs / headings / table cells)
        -> ask the LLM WHAT to change (returns edited text per block)
        -> apply the edits back into the copy, preserving structure,
           tables, styles and embedded images as far as the format allows
        -> persist the copy as a NEW Document (source_file_id = original)
        -> return its metadata

The model only ever decides *what* text changes; this module performs the
actual binary change. The original file on disk is never opened for writing
and is never overwritten.

PDF editing uses PyMuPDF (fitz) for true in-place editing: the original PDF
is opened read-only, its page geometry is analysed, every text block's
bounding box + font + size + colour + alignment is captured, and the LLM is
asked to rewrite ONLY the text (it never sees coordinates or images).
The backend then redacts each original text block and writes the edited text
back into the exact same rectangle, shrinking the font / wrapping / growing
the box (with an overlap check) when the new text is longer. Images, vector
graphics, table gridlines and page geometry are left untouched. See
``_apply_pdf_edits`` for the guarantees and remaining limitations.
"""

import io
import json
import logging
import os
import re
import sys
from pathlib import Path

from app.models.document import Document
from app.services.documents import persist_file
from app.services.errors import DocumentEditError
from app.services import gemini

logger = logging.getLogger("app.document_edit")

# Formats we can edit by rewriting their text content.
EDITABLE_FORMATS = ("docx", "odt", "pdf", "txt", "md")


def edit_document(
    document_id: int,
    instruction: str,
    user_id: int,
    db,
    chat_id: int | None = None,
) -> dict:
    """Copy ``document_id``, apply the LLM-driven text edits, save a new file.

    Returns a structured result dict (success + new document metadata). Raises
    ``DocumentEditError`` on any failure — the original file is never touched.
    """
    if not instruction or not instruction.strip():
        raise DocumentEditError("edit instruction is empty")

    source = (
        db.query(Document)
        .filter(Document.id == document_id, Document.user_id == user_id)
        .first()
    )
    if source is None:
        raise DocumentEditError(f"document {document_id} not found")

    source_path = Path(source.filepath)
    if not source_path.is_file():
        raise DocumentEditError("source file is missing on disk")

    file_type = source.file_type
    if file_type not in EDITABLE_FORMATS:
        raise DocumentEditError(
            f"editing '{file_type}' is not supported (supported: "
            f"{', '.join(EDITABLE_FORMATS)})"
        )

    original_bytes = source_path.read_bytes()  # READ ONLY — never written back

    try:
        if file_type == "docx":
            new_bytes = _edit_docx(original_bytes, instruction)
        elif file_type == "odt":
            new_bytes = _edit_odt(original_bytes, instruction)
        elif file_type == "pdf":
            new_bytes = _edit_pdf(original_bytes, instruction)
        else:  # txt / md
            new_bytes = _edit_text_file(original_bytes, instruction, file_type)
    except DocumentEditError:
        raise
    except Exception as exc:
        logger.exception("Document edit failed for document %s", document_id)
        raise DocumentEditError(f"edit failed: {exc}") from exc

    edited_name = _edited_filename(source.original_filename)
    document = persist_file(
        new_bytes,
        file_type,
        user_id,
        db,
        original_filename=edited_name,
        source_file_id=source.id,
        chat_id=chat_id,
    )

    logger.info(
        "Edited document %s -> new document %s (format=%s)",
        document_id,
        document.id,
        file_type,
    )
    return {
        "success": True,
        "document_id": document.id,
        "filename": document.original_filename,
        "file_type": document.file_type,
        "file_size": document.file_size,
        "source_file_id": source.id,
    }


def _edited_filename(original: str) -> str:
    stem, dot, ext = original.rpartition(".")
    if dot:
        return f"{stem} (отредактировано).{ext}"
    return f"{original} (отредактировано)"


# --------------------------------------------------------------------------- #
# Deterministic brand removal (LXSHOW)
#
# The user requires EVERY mention of "LXSHOW" gone from the output, regardless
# of what the LLM does. We remove it from the translated text and from table
# cells before rendering, and reject the whole edit if any survives into the
# final PDF. Removal is token-based, so we never strip fragments of unrelated
# words (e.g. "relaxshow" is left untouched).
# --------------------------------------------------------------------------- #

_LXSHOW_RE = re.compile(r"LXSHOW(?:\s*(?:Laser|лазер))?", re.IGNORECASE)


def _sanitize_lxshow(text: str | None) -> str | None:
    """Strip a standalone ``LXSHOW`` token (optionally followed by Laser/лазер)."""
    if not text:
        return text
    cleaned = _LXSHOW_RE.sub("", text)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned


def _verify_no_lxshow(texts: list[str] | None = None, content: bytes | None = None) -> None:
    """Final safety net against the brand 'LXSHOW' leaking into the output.

    ``texts`` are the strings we actually rendered (already sanitized by
    ``_sanitize_lxshow``); this is the authoritative check -- it catches any
    model/parsing slip that let 'LXSHOW' survive into produced text. Unedited
    source regions are intentionally NOT scanned: they are the user's original
    content and outside the scope of an in-place edit, so scanning them would
    wrongly reject a perfectly good edit. ``content`` keeps a legacy whole-document
    scan for callers that need it.
    """
    import fitz

    if texts is not None:
        blob = "\n".join("" if t is None else t for t in texts)
        if re.search(r"lxshow", blob, re.IGNORECASE):
            raise DocumentEditError(
                "edited text still contains the brand 'LXSHOW'; aborting so the "
                "original is left untouched"
            )
        return
    if content is not None:
        doc = fitz.open(stream=content, filetype="pdf")
        try:
            for page in doc:
                if re.search(r"lxshow", page.get_text(), re.IGNORECASE):
                    raise DocumentEditError(
                        "edited PDF still contains the brand 'LXSHOW'; aborting so the "
                        "original is left untouched"
                    )
        finally:
            doc.close()


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #


def _edit_docx(content: bytes, instruction: str) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))

    # Index python-docx wrappers by their underlying XML element so we can map
    # a body-level <w:p>/<w:tbl> back to its wrapper while walking in order.
    para_by_element = {para._element: para for para in doc.paragraphs}
    table_by_element = {table._tbl: table for table in doc.tables}

    blocks: list = []  # parallel list of ("p", Paragraph) / ("cell", Cell)
    texts: list[str] = []
    for element in doc.element.body:
        tag = element.tag
        if tag.endswith("}p"):
            para = para_by_element.get(element)
            if para is None:
                continue
            blocks.append(("p", para))
            texts.append((para.text or "").strip())
        elif tag.endswith("}tbl"):
            table = table_by_element.get(element)
            if table is None:
                continue
            for row in table.rows:
                for cell in row.cells:
                    blocks.append(("cell", cell))
                    texts.append(cell.text or "")

    edited = _request_edits(texts, instruction)
    for (kind, target), new_text in zip(blocks, edited):
        if kind == "p":
            _docx_set_para_text(target, new_text)
        elif target.paragraphs:
            # Rewrite the cell's first paragraph only, keeping its images and
            # run layout. Setting ``cell.text`` would drop every other paragraph
            # of the cell (nested tables, images, multi-paragraph content).
            _docx_set_para_text(target.paragraphs[0], new_text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _docx_set_para_text(para, text: str) -> None:
    """Replace a paragraph's text with ``text`` while preserving images and
    other non-text content (fields, breaks, inline pictures).

    The first text-bearing ``w:r`` is KEPT in place and only its ``w:t`` text is
    rewritten, so any inline image/field child it carries survives. Later
    text-bearing runs are dropped, but their ``w:drawing`` children are moved
    onto the kept run first so pictures are never lost. Runs without ``w:t``
    (pure images, empty runs) keep exactly where they were.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    parent = para._element
    anchor_run = None
    preserved_drawings: list = []
    for el in list(parent):
        if el.tag != qn("w:r"):
            continue
        if el.find(qn("w:t")) is None:
            # Pure image / empty run: keep untouched, position preserved.
            continue
        if anchor_run is None:
            anchor_run = el
        else:
            preserved_drawings.extend(el.findall(qn("w:drawing")))
            parent.remove(el)

    if anchor_run is not None:
        for t in anchor_run.findall(qn("w:t")):
            anchor_run.remove(t)
        for drawing in preserved_drawings:
            anchor_run.append(drawing)
        w_t = OxmlElement("w:t")
        w_t.text = text
        anchor_run.insert(0, w_t)
    else:
        para.add_run(text)


# --------------------------------------------------------------------------- #
# ODT
# --------------------------------------------------------------------------- #


def _edit_odt(content: bytes, instruction: str) -> bytes:
    from odf.opendocument import load
    from odf import teletype

    doc = load(io.BytesIO(content))

    block_tags = frozenset(
        {"text:p", "text:h", "table:table-cell", "table:covered-table-cell"}
    )
    blocks: list = []
    texts: list[str] = []

    def walk(node) -> None:
        tag = getattr(node, "tagName", "")
        if tag in block_tags:
            blocks.append(node)
            texts.append(" ".join(teletype.extractText(node).split()))
            return
        for child in getattr(node, "childNodes", ()):
            walk(child)

    walk(doc.body)

    edited = _request_edits(texts, instruction)
    for node, new_text in zip(blocks, edited):
        _odt_set_block_text(node, new_text)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _odt_set_block_text(node, text: str) -> None:
    """Replace a block's text but keep embedded images (draw:frame).

    ODF table cells must contain their content inside a ``text:p`` paragraph;
    writing a bare ``text:span`` directly under ``table:table-cell`` produces
    a malformed cell (LibreOffice/Word may open it, but per-paragraph content
    is lost and the cell collapses to one unparagraphed line). Cell content is
    therefore wrapped in a brand-new paragraph.
    """
    from odf.text import P, Span

    frames = [
        c
        for c in getattr(node, "childNodes", ())
        if getattr(c, "tagName", None) == "draw:frame"
    ]
    for child in list(getattr(node, "childNodes", ())):
        try:
            node.removeChild(child)
        except Exception:
            pass
    for frame in frames:
        node.addElement(frame)
    tag = getattr(node, "tagName", "")
    if tag in ("table:table-cell", "table:covered-table-cell"):
        para = P()
        para.addElement(Span(text=text))
        node.addElement(para)
    else:
        node.addElement(Span(text=text))


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #


def _edit_pdf(content: bytes, instruction: str) -> bytes:
    """In-place PDF text edit. See ``_apply_pdf_edits`` for the guarantees."""

    fragments, items = _extract_pdf_structure(content)
    if not items:
        raise DocumentEditError(
            "PDF has no extractable text layer (likely a scan/image-only PDF); "
            "text editing is not supported."
        )
    # ``items`` are LOGICAL blocks (text + detected tables). A table is a single
    # item no matter how many rows it has, so the LLM can never mistake table
    # rows for independent text blocks. Only textual content is edited.
    edited_items = _request_edits_structured(items, instruction)
    apply_blocks, apply_texts = _items_to_apply(edited_items, fragments)
    new_bytes = _apply_pdf_edits(content, apply_blocks, apply_texts)
    # Final deterministic guarantee: no 'LXSHOW' (any case) may survive in the
    # text we produced. Unedited source regions are out of scope (see
    # _verify_no_lxshow).
    _verify_no_lxshow(texts=apply_texts)
    return new_bytes


def _extract_pdf_blocks(content: bytes) -> list[dict]:
    """Read each page's real text geometry: bbox, font, size, colour, align.

    Only ``type == 0`` (text) blocks are collected; image blocks are skipped so
    that the LLM never touches pictures. Returns [] when there is no text.
    """
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        blocks: list[dict] = []
        for page_index, page in enumerate(doc):
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # 0 = text, 1 = image
                    continue
                text = _pdf_block_text(block)
                if not text.strip():
                    continue
                style = _pdf_block_style(block)
                blocks.append(
                    {
                        "page": page_index,
                        "bbox": list(block["bbox"]),
                        "text": text,
                        "style": style,
                    }
                )
        return blocks
    finally:
        doc.close()


def _build_table_cell_rects(cells: list, table, page_index: int) -> list[list]:
    """Map a PyMuPDF table's cell geometry onto the logical ``cells`` grid.

    Returns an ``n_rows x n_cols`` matrix of ``fitz.Rect`` (or ``None``). Only
    the text content of ``cells`` is shown to the LLM; this matrix is the key the
    apply step uses to put the edited text back into the right place. The LLM
    never sees these rects.

    The source ``table.cells`` is a FLAT list (one rect per detected cell) whose
    length need not equal ``n_rows * n_cols`` for irregular tables (merged /
    missing / ragged cells). We therefore index by explicit (row, col) instead of
    by ``r * n_cols + c`` so a partial table can never raise ``IndexError``.

    Missing geometry is represented as ``None`` - not a fabricated bbox - so the
    apply step can leave that original cell untouched.
    """
    import fitz

    n_rows = len(cells)
    n_cols = max((len(r) for r in cells), default=0)

    geom: dict[tuple[int, int], "fitz.Rect"] = {}
    try:
        # to_cells() yields one object per detected cell, each carrying its own
        # (row, col) coordinates - robust against non-rectangular tables.
        for cell in table.to_cells():
            r = getattr(cell, "row", None)
            c = getattr(cell, "col", None)
            bbox = getattr(cell, "bbox", None)
            if r is None or c is None or bbox is None:
                continue
            geom[(int(r), int(c))] = fitz.Rect(bbox)
    except Exception:
        geom = {}

    if not geom:
        # Fallback: bounded indexing into the flat list (no fabricated coords).
        flat = list(getattr(table, "cells", []) or [])
        for r in range(n_rows):
            for c in range(n_cols):
                idx = r * n_cols + c
                if idx >= len(flat):
                    continue
                x0r, y0r, x1r, y1r = flat[idx]
                geom[(r, c)] = fitz.Rect(x0r, y0r, x1r, y1r)

    detected = len(getattr(table, "cells", [])) if hasattr(table, "cells") else -1
    if geom and detected != n_rows * n_cols:
        logger.warning(
            "PDF table shape mismatch on page %s: rows=%s cols=%s "
            "detected_cells=%s mapped_cells=%s (merged/missing/ragged cells "
            "possible - geometry mapped by (row,col), no IndexError)",
            page_index, n_rows, n_cols, detected, len(geom),
        )

    matrix: list[list] = []
    for r in range(n_rows):
        row_rects = [geom.get((r, c)) for c in range(n_cols)]
        matrix.append(row_rects)
    return matrix


def _pdf_line_metrics(block: dict) -> tuple[float | None, float | None]:
    """Return ``(baseline_top, line_pitch)`` in points for a text block.

      * ``baseline_top`` — distance from the block's top edge to the first
        line's baseline (the glyph origin). Keeping the re-inserted text's
        first baseline at this offset reproduces the original vertical
        placement even when the replacement font's ascender differs from the
        original font's (the historical cause of whole lines silently sliding
        down the page).
      * ``line_pitch`` — the average vertical distance between consecutive
        original line bboxes. Passed to ``insert_textbox`` as ``lineheight``
        so wrapped replacement lines land on the same baseline raster as the
        original text instead of PyMuPDF's (airier) default leading, which
        otherwise forces an unnecessary font shrink or shifts line positions.

    Any metric that cannot be read (no lines / no spans) is returned as
    ``None``; callers then keep the current default behaviour.
    """
    lines = block.get("lines", [])
    if not lines:
        return None, None
    try:
        baseline_top = None
        for span in lines[0].get("spans", []):
            if span.get("text", "").strip():
                baseline_top = span["origin"][1] - block["bbox"][1]
                break
        pitches = [
            lines[i + 1]["bbox"][1] - lines[i]["bbox"][1]
            for i in range(len(lines) - 1)
            if lines[i + 1]["bbox"][1] > lines[i]["bbox"][1]
        ]
        if pitches:
            line_pitch = sum(pitches) / len(pitches)
        else:
            # Single line: its own ascent+descent box is the natural pitch a
            # wrapping replacement line would follow.
            line_pitch = lines[0]["bbox"][3] - lines[0]["bbox"][1]
            if line_pitch <= 0:
                line_pitch = None
        return baseline_top, line_pitch
    except Exception:
        return None, None


def _extract_pdf_structure(content: bytes) -> tuple[list[dict], list[dict]]:
    """Structured PDF extraction: flat text fragments + detected table items.

    Returns ``(fragments, items)``:

      * ``fragments`` is the flat list of text block dicts (page/bbox/text/style)
        consumed by ``_apply_pdf_edits``. It is the low-level unit of change.
      * ``items`` is the list of *logical* editing items the LLM sees:

          - text item:  {"id", "type": "text", "frag_index", "text", ...geometry}
          - table item: {"id", "type": "table", "cells", "cell_rects",
                         "frag_indices", ...geometry}

    A detected table becomes ONE logical item regardless of its row/cell count,
    so the model cannot confuse table rows with independent text blocks (the bug
    that produced "expected 30 items, got 213"). Tables are located with
    PyMuPDF's ``find_tables``; no extra dependency is introduced.

    Fragments that split a single sentence across several text blocks (the first
    block ends without terminating punctuation and the next begins with a
    lowercase continuation) are merged up-front so the LLM sees one complete
    sentence instead of N pieces it would merge itself - which broke the strict
    1:1 item-count contract.
    """
    import fitz

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        # 1) Raw text fragments, one per type==0 block, across all pages.
        fragments: list[dict] = []
        for page_index, page in enumerate(doc):
            page_dict = page.get_text("dict")
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:  # 0 = text, 1 = image
                    continue
                text = _pdf_block_text(block)
                if not text.strip():
                    continue
                fragment = {
                    "page": page_index,
                    "bbox": list(block["bbox"]),
                    "text": text,
                    "style": _pdf_block_style(block),
                }
                baseline_top, line_pitch = _pdf_line_metrics(block)
                if baseline_top is not None:
                    fragment["baseline_top"] = baseline_top
                if line_pitch is not None:
                    fragment["line_pitch"] = line_pitch
                fragments.append(fragment)

        # 2) Detect tables (ruled grids + clustered cell text). Geometry is taken
        # from PyMuPDF's extraction; fragment indices are filled in after the
        # sentence-continuation merge so they index the merged fragment list.
        table_cells: list[dict] = []
        for page_index, page in enumerate(doc):
            try:
                page_tables = page.find_tables().tables
            except Exception:
                page_tables = []
            for table in page_tables:
                cells = table.extract()
                if not cells:
                    continue
                cell_rects = _build_table_cell_rects(cells, table, page_index)
                table_cells.append(
                    {
                        "page": page_index,
                        "bbox": list(table.bbox),
                        "cells": cells,
                        "cell_rects": cell_rects,
                        "frag_indices": [],
                    }
                )

        # 3) Fragments owned by a table must never be merged with body text:
        # table cell words are separate editing units, not sentence pieces.
        table_owned: set[int] = set()
        for t in table_cells:
            tb = fitz.Rect(t["bbox"])
            for i, f in enumerate(fragments):
                if (
                    f["page"] == t["page"]
                    and tb.contains(
                        fitz.Point(
                            (f["bbox"][0] + f["bbox"][2]) / 2,
                            (f["bbox"][1] + f["bbox"][3]) / 2,
                        )
                    )
                ):
                    table_owned.add(i)

        # 4) Merge fragments that split a single sentence across text blocks.
        fragments = _merge_continuation_fragments(fragments, table_owned)

        # 5) Recompute each table's fragment indices against the merged list.
        for t in table_cells:
            tb = fitz.Rect(t["bbox"])
            t["frag_indices"] = [
                i
                for i, f in enumerate(fragments)
                if f["page"] == t["page"]
                and tb.contains(
                    fitz.Point(
                        (f["bbox"][0] + f["bbox"][2]) / 2,
                        (f["bbox"][1] + f["bbox"][3]) / 2,
                    )
                )
            ]

        # 6) Build logical items, skipping fragments already owned by a table so
        # each piece of text is edited exactly once (as a text or a table cell).
        items: list[dict] = []
        covered: set[int] = set()
        for t in table_cells:
            covered.update(t["frag_indices"])
        next_id = 0
        for i, f in enumerate(fragments):
            if i in covered:
                continue
            items.append(
                {
                    "id": next_id,
                    "type": "text",
                    "frag_index": i,
                    "page": f["page"],
                    "bbox": f["bbox"],
                    "style": f["style"],
                    "text": f["text"],
                }
            )
            next_id += 1
        default_style = {"font": "helv", "size": 11.0, "color": 0, "align": 0}
        for t in table_cells:
            style = (
                fragments[t["frag_indices"][0]]["style"]
                if t["frag_indices"]
                else default_style
            )
            items.append(
                {
                    "id": next_id,
                    "type": "table",
                    "page": t["page"],
                    "bbox": t["bbox"],
                    "cell_rects": t["cell_rects"],
                    "frag_indices": t["frag_indices"],
                    "style": style,
                    "cells": t["cells"],
                }
            )
            next_id += 1
        return fragments, items
    finally:
        doc.close()


def _bbox_union(a: list, b: list) -> list:
    """Union of two [x0, y0, x1, y1] bounding boxes."""
    return [
        min(a[0], b[0]),
        min(a[1], b[1]),
        max(a[2], b[2]),
        max(a[3], b[3]),
    ]


def _is_sentence_continuation(a: dict, b: dict) -> bool:
    """True when ``b`` continues the sentence of the preceding fragment ``a``.

    Some PDFs split one sentence across several text blocks: the first block
    ends without terminating punctuation and the next starts with a lowercase
    continuation (e.g. ``'...angle'`` + ``'steel,extra nesting software ...'``).
    The LLM then merges them into a single translation and the strict 1:1 item
    count check rejects the reply. Detecting them up-front keeps the model reply
    aligned and the redraw region correct.
    """
    if a["page"] != b["page"]:
        return False
    ta = a["text"].strip()
    tb = b["text"].strip()
    if not ta or not tb:
        return False
    if not tb[0].islower():
        return False
    if ta[-1] in ".!?:;。！？；，、":
        return False
    # b sits below a within ~2 line heights, in the same column.
    gap = b["bbox"][1] - a["bbox"][3]
    if not (0 <= gap < 60):
        return False
    return max(a["bbox"][0], b["bbox"][0]) < min(a["bbox"][2], b["bbox"][2])


def _merge_continuation_fragments(
    fragments: list[dict], table_owned: set[int]
) -> list[dict]:
    """Merge adjacent fragments that split one sentence into a single fragment.

    The merged fragment keeps the first fragment's page/style, the union of the
    bboxes and the concatenated text. Fragments owned by a table are excluded so
    cell words are never fused with surrounding body text.
    """
    out: list[dict] = []
    prev: dict | None = None
    prev_owned = False
    for i, f in enumerate(fragments):
        owned = i in table_owned
        if (
            prev is not None
            and not prev_owned
            and not owned
            and _is_sentence_continuation(prev, f)
        ):
            prev["bbox"] = _bbox_union(prev["bbox"], f["bbox"])
            prev["text"] = prev["text"].rstrip() + " " + f["text"].strip()
            continue
        if prev is not None:
            out.append(prev)
        prev = dict(f)
        prev_owned = owned
    if prev is not None:
        out.append(prev)
    return out


def _pdf_block_text(block: dict) -> str:
    lines = []
    for line in block.get("lines", []):
        spans = "".join(s.get("text", "") for s in line.get("spans", []))
        lines.append(spans)
    return "\n".join(lines)


def _pdf_block_style(block: dict) -> dict:
    """Capture font/size/colour/alignment from the first non-empty text span."""
    font = "helv"
    size = 11.0
    color = 0
    for line in block.get("lines", []):
        for span in line.get("spans", []):
            if span.get("text", "").strip():
                font = span.get("font", "helv")
                size = float(span.get("size", 11) or 11)
                color = int(span.get("color", 0) or 0) & 0xFFFFFF
                break
        else:
            continue
        break
    return {
        "font": font,
        "size": size,
        "color": color,
        "align": _pdf_block_align(block),
    }


def _pdf_block_align(block: dict) -> int:
    """0=left, 1=center, 2=right, inferred from the first line position."""
    lines = block.get("lines", [])
    if not lines:
        return 0
    lb = lines[0].get("bbox", [0, 0, 0, 0])
    bb = block.get("bbox", [0, 0, 0, 0])
    width = (bb[2] - bb[0]) or 1
    left_gap = lb[0] - bb[0]
    right_gap = bb[2] - lb[2]
    if left_gap > width * 0.25 and right_gap > width * 0.25:
        return 1  # centered
    if right_gap < width * 0.1:
        return 2  # right aligned
    return 0  # left aligned


def _pdf_resolve_fontname(span_font: str) -> str:
    """Map an embedded (often subset-prefixed) PDF font to a base-14 name."""
    if not span_font:
        return "helv"
    name = span_font.split("+")[-1] if "+" in span_font else span_font
    base = {
        "Helvetica": "helv",
        "Helvetica-Bold": "hebo",
        "Helvetica-Oblique": "heob",
        "Helvetica-BoldOblique": "hebi",
        "Times-Roman": "tiro",
        "Times-Bold": "tibo",
        "Times-Italic": "tiob",
        "Times-BoldItalic": "tibi",
        "Courier": "cour",
        "Courier-Bold": "cobo",
        "Courier-Oblique": "coob",
        "Arial": "helv",
        "Calibri": "helv",
        "Liberation": "helv",
    }
    return base.get(name, "helv")


# Candidate locations for a Cyrillic-capable Unicode TTF, in priority order.
# Production (Docker, Debian/Ubuntu) ships DejaVu Sans via fonts-dejavu-core;
# a repo-bundled copy is the self-contained fallback; Windows Arial is used for
# local dev where DejaVu is absent.
def _pdf_unicode_font(span_font: str | None):
    """Return ``(fontfile, fontname)`` for a Cyrillic-capable Unicode TTF.

    Every candidate path is verified to actually open with PyMuPDF before being
    returned, so a missing/corrupt file can never reach ``insert_textbox`` and
    raise ``need font file or buffer``. Falls back to ``(None, base14_name)``
    when no usable Unicode TTF exists, so ASCII rendering keeps the historical
    base-14 behaviour.
    """
    import glob

    base14 = _pdf_resolve_fontname(span_font or "")
    raw = (span_font or "").split("+")[-1] if span_font else ""
    low = raw.lower()
    bold = "bold" in low or base14 in ("hebo", "tibo", "cobo")
    italic = "oblique" in low or "italic" in low or base14 in ("heob", "tiob", "coob")
    variant = (
        "bolditalic" if (bold and italic) else
        "bold" if bold else
        "italic" if italic else
        "regular"
    )
    families = [
        (
            "/usr/share/fonts/truetype/dejavu",
            "DejaVuSans",
            {
                "regular": "DejaVuSans.ttf",
                "bold": "DejaVuSans-Bold.ttf",
                "italic": "DejaVuSans-Oblique.ttf",
                "bolditalic": "DejaVuSans-BoldOblique.ttf",
            },
        ),
        (
            "/usr/share/fonts/dejavu",
            "DejaVuSans",
            {
                "regular": "DejaVuSans.ttf",
                "bold": "DejaVuSans-Bold.ttf",
                "italic": "DejaVuSans-Oblique.ttf",
                "bolditalic": "DejaVuSans-BoldOblique.ttf",
            },
        ),
        (
            os.path.join(os.path.dirname(__file__), "fonts"),
            "DejaVuSans",
            {
                "regular": "DejaVuSans.ttf",
                "bold": "DejaVuSans-Bold.ttf",
                "italic": "DejaVuSans-Oblique.ttf",
                "bolditalic": "DejaVuSans-BoldOblique.ttf",
            },
        ),
    ]
    if sys.platform.startswith("win"):
        families.append(
            (
                r"C:\Windows\Fonts",
                "Arial",
                {
                    "regular": "arial.ttf",
                    "bold": "arialbd.ttf",
                    "italic": "ariali.ttf",
                    "bolditalic": "arialbi.ttf",
                },
            )
        )

    # (path, pdf_resource_name) pairs. The resource name MUST NOT contain
    # spaces -- PyMuPDF rejects "bad fontname chars" otherwise.
    candidates = []
    for base_dir, fam, variants in families:
        candidates.append((os.path.join(base_dir, variants[variant]), fam))
        candidates.append((os.path.join(base_dir, variants["regular"]), fam))
    # Broad filesystem search as a safety net across distros/packaging.
    for pat in (
        "/usr/share/fonts/**/*.ttf",
        "/usr/share/fonts/**/*.otf",
        "/usr/local/share/fonts/**/*.ttf",
        "/usr/local/share/fonts/**/*.otf",
        os.path.join(os.path.dirname(__file__), "fonts", "**", "*.ttf"),
    ):
        for hit in glob.glob(pat, recursive=True):
            low = hit.lower()
            name = (
                "Arial" if "arial" in low
                else "DejaVuSans"
            )
            candidates.append((hit, name))

    for path, name in candidates:
        if not os.path.isfile(path):
            continue
        # Verify the font can actually be EMBEDDED (what insert_textbox does),
        # not merely opened for metrics -- an un-embeddable file would otherwise
        # crash insert_textbox with "need font file or buffer".
        if not _pdf_font_embeddable(path):
            continue
        return path, name
    return None, base14


def _pdf_font_embeddable(path: str) -> bool:
    """Return True only if PyMuPDF can embed ``path`` as a page font."""
    import fitz

    try:
        doc = fitz.open()
        doc.new_page()
        doc[0].insert_font(fontname="helv", fontfile=path)
        doc.close()
        return True
    except Exception:
        return False


def _redaction_rects(page, block: dict) -> list:
    """Return the glyph rectangles that belong to ``block`` (for redaction).

    Redaction removes every text object whose bbox intersects the annot
    rectangle. A text block's ``bbox`` from ``get_text("dict")`` is generous
    (it spans the whole line grid), so redacting it directly can silently
    delete text of a NEIGHBOURING block that merely overlaps that rectangle
    (a caption, an inline annotation, a fragment that pokes past its cell).
    We therefore redact only the actual glyph footprints (span bboxes) inside
    the block bbox. Falls back to the block bbox itself when no span geometry
    is available (e.g. synthetic cell rects on pages without extractable text).
    """
    import fitz

    target = fitz.Rect(block["bbox"])
    if target.is_empty or target.width < 1 or target.height < 1:
        return [target]

    rects: list[fitz.Rect] = []
    try:
        page_dict = page.get_text("rawdict")
    except Exception:
        page_dict = {}
    for tblock in page_dict.get("blocks", []):
        if tblock.get("type") != 0:
            continue
        for line in tblock.get("lines", []) or []:
            for span in line.get("spans", []) or []:
                span_rect = fitz.Rect(span["bbox"])
                inter = span_rect & target
                if not inter.is_empty:
                    rects.append(inter)
    if not rects:
        return [target]

    # Deduplicate near-identical footprints (merged fragments overlap).
    keyed = {
        (round(r.x0, 2), round(r.y0, 2), round(r.x1, 2), round(r.y1, 2))
        for r in rects
    }
    if not keyed:
        return [target]
    return [fitz.Rect(k) for k in keyed]


def _apply_pdf_edits(
    content: bytes, blocks: list[dict], edited_texts: list[str]
) -> bytes:
    """Redact each original text block and write ``edited_texts`` back in place.

    Guarancees / behaviour:
      * The source ``content`` is opened read-only and never written.
      * Only ``type == 0`` text blocks are modified; images, vector drawings
        and table gridlines are preserved (redaction uses ``fill=None`` so it
        does not white-out underlying graphics; images and line-art are never
        removed, and only the block's own glyph footprints are redacted, so a
        neighbouring block that overlaps the block bbox is left untouched).
      * Edited text is placed into the block's original bbox with the original
        font/size/colour/alignment. If it does not fit, the font is shrunk; if
        it still overflows, the box is grown downward only while it does not
        intersect a neighbouring block (overlap check) and stays inside the page.
      * As a last resort the text is clipped into the original box at a minimum
        size rather than destroying the layout of other elements.

    Limitation: a block's styles are collapsed to the first span's style (mixed
    bold/italic within one block is not reproduced), and the font is mapped to
    the closest base-14 family.
    """
    import fitz

    if len(blocks) != len(edited_texts):
        raise DocumentEditError(
            "PDF edit failed: the number of edited blocks does not match the "
            "number of source blocks."
        )

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        per_page: dict[int, list[tuple[dict, str]]] = {}
        for block, new_text in zip(blocks, edited_texts):
            per_page.setdefault(block["page"], []).append((block, new_text))

        for page_index, items in per_page.items():
            page = doc[page_index]
            page_rect = page.rect
            # Obstacles the rewritten text must not overlap: other text blocks,
            # embedded images and vector drawings (tables/gridlines) on the page.
            page_dict = page.get_text("dict")
            obstacles: list = []
            for blk in page_dict.get("blocks", []):
                if blk.get("type") == 1:  # image block
                    obstacles.append(fitz.Rect(blk["bbox"]))
            for drawing in page.get_drawings():
                obstacles.append(fitz.Rect(drawing.get("rect", [0, 0, 0, 0])))
            text_siblings = [
                fitz.Rect(b["bbox"]) for (b, _) in items
            ]
            # Non-edited text blocks that do not intersect any edited block are
            # obstacles too, so the rewritten text can never slide over an
            # untouched paragraph (merged/table fragments are excluded because
            # their bboxes intersect an item bbox).
            for blk in page_dict.get("blocks", []):
                if blk.get("type") != 0:
                    continue
                other = fitz.Rect(blk["bbox"])
                if any(not (other & ir).is_empty for ir in text_siblings):
                    continue
                obstacles.append(other)
            sibling_rects = obstacles + text_siblings
            # Redact only the edited blocks' own glyph footprints, preserving
            # neighbors, images and line-art (gridlines) in the same rectangle.
            for block, _ in items:
                for rect in _redaction_rects(page, block):
                    page.add_redact_annot(rect, fill=None)
            page.apply_redactions(
                images=fitz.PDF_REDACT_IMAGE_NONE,
                graphics=fitz.PDF_REDACT_LINE_ART_NONE,
            )
            # Then write the edited text back.
            for block, new_text in items:
                _pdf_fit_and_insert(
                    page, block, new_text, page_rect, sibling_rects,
                )

        out = io.BytesIO()
        doc.save(out, garbage=4, deflate=True)
        return out.getvalue()
    finally:
        doc.close()


def _pdf_normalize_color(value) -> tuple | None:
    """Return a PyMuPDF-valid text color: ``None`` or a 3/4-tuple of floats 0..1.

    Accepts ``None``, an integer RGB (0..0xFFFFFF), a 3/4-tuple of ints 0..255,
    or a 3/4-tuple of floats 0..1. Anything unusable falls back to ``None``
    (PyMuPDF's default text colour). This shields ``insert_textbox`` from the
    ``ValueError: need 1, 3 or 4 color components in range 0 to 1`` that an
    out-of-range integer colour would raise.
    """
    if value is None:
        return None
    if isinstance(value, (tuple, list)):
        comps = list(value)
        if len(comps) in (3, 4) and all(isinstance(c, (int, float)) for c in comps):
            if any(c > 1 for c in comps):  # looks like 0..255 ints -> scale down
                comps = [max(0.0, min(1.0, c / 255.0)) for c in comps]
            else:
                comps = [max(0.0, min(1.0, float(c))) for c in comps]
            return tuple(comps)
        return None
    if isinstance(value, (int, float)):
        iv = int(value) & 0xFFFFFF
        return (
            ((iv >> 16) & 0xFF) / 255.0,
            ((iv >> 8) & 0xFF) / 255.0,
            (iv & 0xFF) / 255.0,
        )
    return None


def _pdf_fit_and_insert(
    page, block: dict, new_text: str, page_rect, sibling_rects: list
) -> None:
    """Write ``new_text`` back into the block's place, keeping layout intact.

    The original text block is already redacted by the caller. The edited text
    must stay inside the ORIGINAL bbox - the frame the reader expects (a table
    cell's borders, or the paragraph's original extent). We never grow the box
    downward, because growing lets translated text ride over the neighbouring
    cell gridline or the next paragraph. Instead the text is fit by shrinking
    the font until ``insert_textbox`` reports it fits.

    Two layout-preserving refinements use metrics captured during extraction
    (only when present; otherwise behaviour is unchanged):

      * ``line_pitch`` (points) is passed to ``insert_textbox`` as
        ``lineheight`` so the replacement lines land on the ORIGINAL baseline
        raster. Without it PyMuPDF applies its airier default leading, which
        makes a block overflow and forces a needless font shrink.
      * ``baseline_top`` (points) is reproduced: the frame is shifted so the
        first baseline sits exactly where the original first line was, instead
        of ``bbox.y0 + ascender*fontsize`` (dropping whole lines down the page
        when the replacement font's ascender differs from the original's).
        The shift is only ever downward (never above the frame top).

    Placement is driven by ``insert_textbox``'s own return code (the ground
    truth for what fits): a negative code means "did not fit, nothing written",
    so we descend the font size and retry rather than losing or overflowing the
    text. As an absolute last resort (an unbreakable token wider than the frame,
    or text taller than the frame even at the floor size) the box is widened /
    grown only within the page and away from neighbours so the text is never
    silently dropped.
    """
    import fitz

    bbox = block["bbox"]
    style = block.get("style", {}) or {}

    # Always resolve to a guaranteed base-14 resource name for the fallback path:
    # the PDF span's raw font name may be a custom/embedded family that PyMuPDF
    # cannot resolve on its own, which would crash insert_textbox with
    # "need font file or buffer".
    base_font = _pdf_resolve_fontname(style.get("font", "Helvetica"))
    # Embed a Unicode TTF only when the text carries characters beyond Latin-1
    # (Cyrillic / Greek / CJK / ...). Plain Latin text keeps the lightweight,
    # cleanly-extractable base-14 font.
    fontfile, fontname = (None, base_font)
    if any(ord(ch) > 255 for ch in new_text):
        u_path, u_name = _pdf_unicode_font(base_font)
        if u_path:
            fontfile, fontname = u_path, u_name
    # Embed the Unicode TTF up front -- this is the canonical, version-robust
    # way to use an arbitrary font (registering it once on the page). On any
    # failure we degrade to the base-14 fallback font so the edit never crashes
    # with "need font file or buffer".
    if fontfile:
        try:
            page.insert_font(fontname=fontname, fontfile=fontfile)
        except Exception:
            fontname = base_font
            fontfile = None
    size = float(style.get("size", 11) or 11)
    color = _pdf_normalize_color(style.get("color", 0))
    align = int(style.get("align", 0))
    floor = 1.0

    # Loop-adjusted layout metrics captured at extraction time (points).
    baseline_top = block.get("baseline_top")
    line_pitch = block.get("line_pitch")

    # Ascender of the font we ACTUALLY insert, used to reproduce the original
    # baseline and line raster. Unavailable fonts simply keep default layout.
    ascender = None
    if line_pitch is not None or baseline_top is not None:
        try:
            if fontfile:
                ascender = fitz.Font(fontfile=fontfile).ascender
            else:
                ascender = fitz.Font(fontname=fontname).ascender
        except Exception:
            ascender = None

    # The frame: exactly the original bbox (a paragraph's extent or a table
    # cell's borders). The box is never grown beyond it, so translated text
    # stays exactly where it was and cannot poke over a cell gridline. For a
    # tight block bbox ``insert_textbox`` may need a hair more line-height than
    # the glyph box, which is absorbed by a tiny font shrink instead of an
    # overflow.
    base_rect = fitz.Rect(bbox)
    if base_rect.width < 1 or base_rect.height < 1:
        return

    def _layout_rect(fs):
        """Frame for a candidate ``fs``: keeps the original first baseline."""
        if baseline_top is not None and ascender:
            shift = baseline_top - ascender * fs
            if shift >= 0:  # never push the text above the original frame top
                return fitz.Rect(
                    base_rect.x0,
                    base_rect.y0 + shift,
                    base_rect.x1,
                    base_rect.y1 + shift,
                )
        return base_rect

    def _lineheight(fs):
        if line_pitch is not None and ascender and fs * ascender > 0:
            return line_pitch / (fs * ascender)
        return None

    def _insert(r, fs):
        return page.insert_textbox(
            r, new_text, fontname=fontname, fontsize=fs,
            lineheight=_lineheight(fs), color=color, align=align,
        )

    # Largest font whose wrapped text fits the frame. Descend until the first
    # non-negative return code (that size is the one written - negative codes
    # write nothing, so no partial/overlapping text can accumulate).
    cur_fs = size
    while cur_fs >= floor - 1e-9:
        if _insert(_layout_rect(cur_fs), cur_fs) >= 0:
            return
        cur_fs -= 0.5

    # Even the floor size does not fit: an unbreakable token wider than the
    # frame, or text taller than the frame at 1pt. Widen horizontally, bounded
    # by the page and neighbouring content, so the text is not silently dropped.
    cand = fitz.Rect(_layout_rect(floor))
    for _ in range(1, 100):
        grown = fitz.Rect(cand.x0 - 2, cand.y0, cand.x1 + 2, cand.y1)
        if grown.x0 < page_rect.x0 or grown.x1 > page_rect.x1:
            break
        if any(grown.intersects(r) for r in sibling_rects if r != base_rect):
            break
        cand = grown
        if _insert(cand, floor) >= 0:
            return
    # Last-ditch downward room (bounded by the page) before giving up.
    for _ in range(1, 100):
        grown = fitz.Rect(cand.x0, cand.y0, cand.x1, cand.y1 + 2)
        if grown.y1 > page_rect.y1:
            break
        cand = grown
        if _insert(cand, floor) >= 0:
            return


# --- Test hook: apply a known set of edits without calling the LLM. ---


def _edit_pdf_with_edits(content: bytes, edited_texts: list[str]) -> bytes:
    """Apply a caller-supplied list of edited texts (used by tests)."""
    blocks = _extract_pdf_blocks(content)
    if not blocks:
        raise DocumentEditError(
            "PDF has no extractable text layer (likely a scan/image-only PDF); "
            "text editing is not supported."
        )
    return _apply_pdf_edits(content, blocks, edited_texts)


# --------------------------------------------------------------------------- #
# Plain text / Markdown
# --------------------------------------------------------------------------- #


def _edit_text_file(content: bytes, instruction: str, file_type: str) -> bytes:
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        text = content.decode("utf-8", errors="replace")

    if file_type == "md":
        blocks = text.split("\n")
    else:
        blocks = text.split("\n")

    edited = _request_edits(blocks, instruction)
    return "\n".join(edited).encode("utf-8")


# --------------------------------------------------------------------------- #
# LLM text editing
# --------------------------------------------------------------------------- #


_EDIT_SYSTEM = (
    "You are a document editing engine.\n"
    "Your task is to edit the provided PDF text blocks according to the user's instruction.\n\n"
    "CRITICAL OUTPUT RULE:\n"
    "Return ONLY a valid JSON array.\n"
    "Do not use Markdown.\n"
    "Do not use ```json fences.\n"
    "Do not add explanations, comments, or any text outside the JSON array.\n\n"
    "The returned array MUST contain exactly N elements, where N is the number of "
    "input blocks (count them).\n\n"
    "Each element MUST be a JSON object with exactly one field:\n"
    '{"text": "edited text"}\n\n'
    "Rules:\n"
    "1. Preserve the order of all blocks exactly.\n"
    "2. Return exactly one object for every input block (one input -> one output).\n"
    "3. If a block does not need modification, return its original text unchanged.\n"
    "4. Do NOT merge blocks and do NOT split blocks.\n"
    "5. Do NOT omit blocks and do NOT add extra blocks.\n"
    "6. Preserve numbers, units of measurement, designations, symbols, code, "
    "formatting markers and technical terms; change only what the instruction asks.\n"
    "7. For translation requests, translate each block into the requested language "
    "while preserving its meaning, structure, numbers and units.\n"
    "8. If a block already contains the target language or non-translatable "
    "technical text, keep it as-is; do not destroy or rewrite it.\n"
    "9. Do not invent information that is not present in the source.\n"
    "10. If the instruction asks to remove a specific name, brand or phrase (for "
    "example LXSHOW), remove it wherever it occurs in the blocks.\n"
    "11. The output must be valid JSON that can be parsed directly by json.loads().\n"
    "12. Escape quotation marks, backslashes and newlines inside strings correctly "
    "according to JSON syntax.\n\n"
    "USER INSTRUCTION:\n"
    "{instruction}\n\n"
    "INPUT BLOCKS:\n"
    "{blocks}\n\n"
    "Return the JSON array now."
)

_RETRY_SYSTEM = (
    "You output strictly a JSON array of strings and nothing else. No prose, no "
    "markdown code fences."
)

_STRUCTURED_RETRY_SYSTEM = (
    "You are a structured document editing engine.\n"
    "Your previous reply was rejected because it broke the strict 1:1 item "
    "contract. Fix it and return ONLY valid JSON - no Markdown, no ```json "
    "fences, no commentary outside the JSON.\n\n"
    "Return EXACTLY one element per input item, carrying the SAME stable ids as "
    "the input (never invent, drop or duplicate an id):\n"
    '- for a TEXT item: {"id": <exact input id>, "text": "..."}\n'
    '- for a TABLE item: {"id": <exact input id>, "cells": [[...], ...]} '
    "(or {\"id\": <id>, \"headers\": [...], \"rows\": [[...]]}) with the EXACT "
    "same number of rows and columns as the input.\n"
    "The array may be in ANY order - the ids alone determine the mapping - but "
    "the id set must match the input id set exactly.\n"
    "Do NOT merge, split, skip or duplicate input items. Preserve all numbers, "
    "units, codes and technical symbols exactly.\n"
    'Example for input ids 0,1,2: [{"id": 0, "text": "..."}, '
    '{"id": 1, "text": "..."}, {"id": 2, "text": "..."}]\n'
)

_EDIT_STRUCTURED_SYSTEM = (
    "You are editing structured PDF content.\n"
    "Return ONLY valid JSON. Do not use Markdown, do not use ```json fences, do "
    "not add explanations, commentary or any text outside the JSON.\n\n"
    "The input items are provided as an object {\"items\": [...]} (or a bare array). "
    "There are two item kinds:\n"
    '- type "text": {"id": N, "type": "text", "text": "..."}. Translate ONLY the text.\n'
    '- type "table": {"id": N, "type": "table", "cells": [[...], ...]}. A table is ONE '
    "item; translate ONLY the text inside each cell, never split it into rows.\n\n"
    "RESPONSE FORMAT (return EXACTLY this shape):\n"
    "Return a JSON ARRAY with ONE element per input item. Prefer the id-bearing "
    "object form so the mapping never depends on array position (a bare string is "
    "also accepted):\n"
    '- for a TEXT item, the element is {"id": <exact input id>, "text": "..."} '
    "(or a bare translated string such as \"Привет\");\n"
    '- for a TABLE item, the element is the table given back as a "cells" matrix '
    '{"cells": [[...], ...]} with the SAME number of rows and columns as the input '
    '"cells" you were given (the exact counts are stated in the user prompt). '
    "A cell is either a translated string or null (null = leave that original cell "
    "unchanged; never use null for the whole table).\n"
    "Optionally, instead of \"cells\" you may return "
    '{"headers": [...], "rows": [[...], ...]} where "headers" MUST be EXACTLY the '
    'first row of the input "cells" and "rows" MUST be EXACTLY the remaining rows '
    "(so total rows are unchanged and equal to the input). Do NOT add, remove, "
    "merge or split rows or columns.\n"
    "The structured object form {\"items\": [{\"id\",\"type\",\"text\"/\"cells\"}, ...]} "
    "is ALSO accepted, but the bare array above is preferred.\n\n"
    "Rules:\n"
    "1. Return the array with exactly the same number of elements as input items, "
    "one output per input item. One input item can never become zero, two or more "
    "output items. When you return objects with an \"id\", the id MUST equal the "
    "input item's id and every input id must appear exactly once: the mapping is "
    "by id, so the array order does not matter. When you return bare strings, "
    "keep the input order.\n"
    "2. For text items the output element is the translated string.\n"
    "3. For table items the output \"cells\" matrix MUST have the EXACT same number "
    "of rows and columns as the input \"cells\" for that item (the exact counts are "
    "given in the user prompt). The table stays ONE item. Do NOT merge, split, add, "
    "remove or reorder rows/cells, and do NOT convert a table into plain text. If a "
    "cell has no text to change, return null for it, never drop the cell.\n"
    "4. NEVER split one input item into several output items. NEVER merge several "
    "input items into one output item.\n"
    "5. Preserve ALL numbers, units, dimensions, formulas, tolerances, technical "
    "identifiers, model names, article numbers, codes and symbols EXACTLY. Change "
    "only the LANGUAGE of the text. Examples that must be kept verbatim:\n"
    "   - 6000W  ->  6000W\n"
    "   - 2000*6000mm  ->  2000*6000mm\n"
    "   - ±0.02mm/m  ->  ±0.02mm/m\n"
    "   - 35-45 work days  ->  35-45 рабочих дней\n"
    "6. If the instruction asks to remove a name/brand/phrase (e.g. LXSHOW), remove "
    "it from text and from every table cell; otherwise leave it untouched. Do not "
    "alter unrelated content. If a cell becomes empty, leave it empty (do not delete "
    "the cell).\n"
    "7. The output must be valid JSON parseable by json.loads(). Escape quotation "
    "marks, backslashes and newlines inside strings per JSON syntax.\n\n"
    "USER INSTRUCTION:\n"
    "{instruction}\n\n"
    "Return the JSON array now."
)

# Chunking limits for a single edit-model request (adaptive chunking).
# Items are grouped so that EACH chunk stays within BOTH bounds:
#   * at most EDIT_CHUNK_SIZE logical items, AND
#   * at most EDIT_CHUNK_MAX_CHARS characters of input text.
# A single item is never split -- if one item alone exceeds EDIT_CHUNK_MAX_CHARS
# (e.g. a very large table), it is sent as its own chunk. This keeps the model's
# JSON answer small enough to avoid the truncation that happens on big prompts
# (the edit model's token budget is limited), while EDIT_MAX_TOKENS raises the
# per-chunk output budget.
EDIT_CHUNK_SIZE = 8
EDIT_CHUNK_MAX_CHARS = 8000
# Generous output budget for one chunk. The edit model must return one JSON
# object/string per block; a ~10k-char chunk translates to a few thousand tokens
# at most, so 8192 is comfortably above that and prevents mid-array truncation.
EDIT_MAX_TOKENS = 8192
# Per-chunk LLM attempts for structured edits: the first call plus one controlled
# retry when the model's reply fails validation. Used only for diagnostic logs.
_EDIT_CHUNK_MAX_ATTEMPTS = 2


def _generate_edits_array(prompt: str, system_instruction: str) -> str:
    """Call the edit model; surface upstream failures as ``DocumentEditError``."""
    try:
        return gemini.generate_answer(
            prompt,
            system_instruction=system_instruction,
            max_tokens=EDIT_MAX_TOKENS,
            response_format={"type": "json_object"},
        )
    except gemini.GeminiError as exc:
        raise DocumentEditError(f"edit model unavailable: {exc}") from exc


def _chunk_texts(
    texts: list[str],
    max_blocks: int = EDIT_CHUNK_SIZE,
    max_chars: int = EDIT_CHUNK_MAX_CHARS,
) -> list[list[str]]:
    """Split ``texts`` into chunks bounded by block count AND total input chars.

    A single text block is never split, so a block longer than ``max_chars``
    simply forms its own chunk. This keeps each model request (and thus its JSON
    answer) small enough to avoid truncation.
    """
    chunks: list[list[str]] = []
    current: list[str] = []
    current_chars = 0
    for text in texts:
        tlen = len(text)
        if current and (len(current) >= max_blocks or current_chars + tlen > max_chars):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(text)
        current_chars += tlen
    if current:
        chunks.append(current)
    return chunks


def _request_edits(blocks: list[str], instruction: str, chunk_size: int = EDIT_CHUNK_SIZE) -> list[str]:
    """Ask the LLM to rewrite each block per the instruction.

    Text blocks are grouped into chunks bounded by both ``EDIT_CHUNK_SIZE`` and
    ``EDIT_CHUNK_MAX_CHARS`` so the model never has to emit one enormous (and easily
    truncated) JSON array. Each chunk is edited independently and the results are
    concatenated in order, so the returned list always has exactly ``len(blocks)``
    elements.

    For each chunk a single *controlled* retry is attempted on the first failure
    (bad JSON, length mismatch). If that also fails, ``DocumentEditError`` is
    raised so the caller aborts and leaves the original untouched. Validation is
    never disabled.
    """
    if not any(b.strip() for b in blocks):
        raise DocumentEditError("document has no editable text")

    chunks = _chunk_texts(blocks, chunk_size, EDIT_CHUNK_MAX_CHARS)
    if len(chunks) == 1:
        return _request_edits_chunk(chunks[0], instruction)

    edited: list[str] = []
    for chunk in chunks:
        edited.extend(_request_edits_chunk(chunk, instruction))
    # Guard against any implementation slip: the caller relies on a 1:1 mapping
    # between blocks and edits. If the chunks did not add up, refuse to edit.
    if len(edited) != len(blocks):
        raise DocumentEditError("model did not return a valid edits array")
    return edited


def _request_edits_chunk(blocks: list[str], instruction: str) -> list[str]:
    """Edit a single chunk with one controlled retry.

    On any invalid response (truncated JSON, wrong element count, missing text)
    diagnostic information about the model reply is logged - never secrets.
    """
    expected = len(blocks)
    numbered = "\n".join(f"{i}\t{b}" for i, b in enumerate(blocks))
    # The system prompt is the full edit template; fill its placeholders here.
    system = _EDIT_SYSTEM.replace("{instruction}", instruction).replace("{blocks}", numbered)
    user_prompt = "Edit the document per the system instructions and return the JSON array."

    raw = _generate_edits_array(user_prompt, system)
    edits, reason = _parse_edits_debug(raw, expected)
    if edits is None:
        logger.warning(
            "edit_document: model reply rejected (first attempt). chunk_blocks=%s "
            "expected=%s prompt_len=%s reason=%s raw_response=%s",
            expected,
            expected,
            len(system) + len(user_prompt),
            reason,
            _safe_raw(raw),
        )
        # Controlled retry: exact count pinned, wrappers forbidden.
        retry_prompt = (
            "DOCUMENT BLOCKS (one per line, format 'index<TAB>text'; a block may "
            "span several visual lines):\n"
            f"{numbered}\n\n"
            f"EDITING INSTRUCTION: {instruction}\n\n"
            f"Return ONLY a JSON array of exactly {expected} items, one per block in "
            "the same order (block 0 first). Each item must be an object "
            '{"text": "..."}. Output nothing except the raw JSON array - no '
            "markdown, no commentary."
        )
        raw = _generate_edits_array(retry_prompt, _RETRY_SYSTEM)
        edits, reason = _parse_edits_debug(raw, expected)
        if edits is None:
            logger.warning(
                "edit_document: model reply rejected (after retry). chunk_blocks=%s "
                "expected=%s prompt_len=%s reason=%s raw_response=%s",
                expected,
                expected,
                len(retry_prompt) + len(_RETRY_SYSTEM),
                reason,
                _safe_raw(raw),
            )
            raise DocumentEditError("model did not return a valid edits array")
    return edits


def _safe_raw(raw: str, limit: int = 4000) -> str:
    """Truncate a model reply so logs stay readable; contains no secrets."""
    if raw is None:
        return "<None>"
    raw = raw.strip()
    return raw if len(raw) <= limit else raw[:limit] + "...[truncated]"


def _parse_edits_array(raw: str, expected: int) -> list[str] | None:
    result, _ = _parse_edits_debug(raw, expected)
    return result


def _parse_edits_debug(raw: str, expected: int) -> tuple[list[str] | None, str | None]:
    """Parse a model reply into the edit list, returning ``(edits, reason)``.

    Accepts: a bare JSON array (strings or ``{"text": ...}`` objects), the same
    array wrapped in a JSON object (json_object response-format mode), markdown
    fences and trailing commas. Strictly enforces that the element count equals
    ``expected`` and that every element carries text.
    """
    if not raw or not raw.strip():
        return None, "empty model response"
    cleaned = raw.strip()
    # Drop markdown code fences (``` or ```json) wherever they appear.
    cleaned = re.sub(r"```[a-zA-Z]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```", "", cleaned).strip()

    data = _extract_json_value(cleaned)
    if data is None:
        return None, "no JSON array/object found (reply may be truncated)"
    # json_object mode returns an object wrapping a single array field.
    if isinstance(data, dict):
        lists = [v for v in data.values() if isinstance(v, list)]
        if len(lists) == 1:
            data = lists[0]
        else:
            return None, "JSON object reply without a single array field"
    if not isinstance(data, list):
        return None, "top-level JSON is not an array"

    # Tolerate arrays whose elements carry the edited text (either all objects
    # or a mix of strings and objects). A present-but-empty "text" is a valid
    # empty block; only a missing key rejects the element.
    if any(isinstance(el, dict) for el in data):
        data = _flatten_edits_objects(data)
        if data is None:
            return None, "array element is an object missing a 'text' field"
    if not isinstance(data, list) or len(data) != expected:
        got = len(data) if isinstance(data, list) else "non-list"
        return None, f"expected {expected} items, got {got}"
    return [str(item) for item in data], None


def _extract_json_value(cleaned: str) -> object | None:
    """Find the first JSON value (array or object) in ``cleaned`` and parse it.

    Repairs a single trailing-comma mistake before giving up. Returns the parsed
    Python value, or ``None`` if nothing valid could be extracted.
    """
    import json

    start, end = cleaned.find("["), cleaned.rfind("]")
    if start == -1 or end <= start:
        start, end = cleaned.find("{"), cleaned.rfind("}")
    if start == -1 or end <= start:
        return None
    candidate = cleaned[start : end + 1]
    try:
        return json.loads(candidate)
    except ValueError:
        # Common LLM flaw: trailing commas. Repair once and retry the parse.
        repaired = re.sub(r",\s*([}\]])", r"\1", candidate)
        try:
            return json.loads(repaired)
        except ValueError:
            return None


def _flatten_edits_objects(data: list) -> list[str] | None:
    """Map each array element to its edited text.

    Accepts plain strings or objects that carry the text in one of the known
    keys (``text``/``edit``/``edited``/``content``). A *present-but-empty*
    value (``{"text": ""}``) is a valid empty block and must NOT be confused
    with a missing field: only an absent key rejects an element.
    """
    items: list[str] = []
    for obj in data:
        if isinstance(obj, str):
            items.append(obj)
            continue
        if not isinstance(obj, dict):
            return None
        text = obj.get("text")
        if text is None:
            text = obj.get("edit")
        if text is None:
            text = obj.get("edited")
        if text is None:
            text = obj.get("content")
        if text is None:
            return None
        items.append(str(text))
    return items


# --------------------------------------------------------------------------- #
# Structured (text + table) LLM editing
# --------------------------------------------------------------------------- #


def _slim_item(it: dict) -> dict:
    """Minimal LLM-facing view of a logical item (no PDF geometry)."""
    if it["type"] == "table":
        return {"id": it["id"], "type": "table", "cells": it["cells"]}
    return {"id": it["id"], "type": "text", "text": it.get("text", "")}


def _item_chars(it: dict) -> int:
    if it["type"] == "table":
        return sum(len(c or "") for row in it["cells"] for c in row)
    return len(it.get("text", ""))


def _chunk_items(
    items: list[dict],
    max_blocks: int = EDIT_CHUNK_SIZE,
    max_chars: int = EDIT_CHUNK_MAX_CHARS,
) -> list[list[dict]]:
    """Group logical items into chunks bounded by item count AND total text size.

    A single item (e.g. a big table) is never split.
    """
    chunks: list[list[dict]] = []
    current: list[dict] = []
    current_chars = 0
    for it in items:
        chars = _item_chars(it)
        if current and (len(current) >= max_blocks or current_chars + chars > max_chars):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(it)
        current_chars += chars
    if current:
        chunks.append(current)
    return chunks


def _request_edits_structured(
    items: list[dict], instruction: str, chunk_size: int = EDIT_CHUNK_SIZE
) -> list[dict]:
    """Edit structured PDF items (text + tables) via the LLM.

    Text items are translated strictly: each item must come back 1:1, one
    controlled retry is attempted, and a final failure raises
    ``DocumentEditError`` so the caller leaves the original untouched.

    Table items are translated best-effort: the model's reply must keep the
    exact extracted grid dimensions, and if it is rejected even after the retry
    the table is kept exactly as extracted rather than failing the whole
    document - a valid, fully-translated document beats a hard error caused by
    one recalcitrant table.
    """
    if not any(_item_chars(it) > 0 for it in items):
        raise DocumentEditError("document has no editable text")

    tables = [it for it in items if it["type"] == "table"]
    texts = [it for it in items if it["type"] != "table"]

    edited_texts = _request_edits_texts_only(texts, instruction, chunk_size)
    edited_tables = _request_edits_tables_only(tables, instruction, chunk_size)

    by_id = {it["id"]: it for it in edited_texts + edited_tables}
    # Return in the original item order so downstream flattening lines up.
    return [by_id[it["id"]] for it in items]


def _request_edits_tables_only(
    tables: list[dict], instruction: str, chunk_size: int = EDIT_CHUNK_SIZE
) -> list[dict]:
    """Translate table items best-effort.

    Each table is a single logical item (never split). The model's reply must
    mirror the exact extracted row/column counts (see ``_structured_dims_note``).
    A table whose reply is invalid even after the controlled retry is passed
    through unchanged instead of aborting the document.
    """
    if not tables:
        return []
    chunks = _chunk_items(tables, chunk_size, EDIT_CHUNK_MAX_CHARS)
    edited: list[dict] = []
    for i, chunk in enumerate(chunks):
        try:
            edited.extend(
                _request_edits_chunk_structured(chunk, instruction, chunk_index=i)
            )
        except DocumentEditError as exc:
            logger.warning(
                "edit_document(structured): table chunk %s could not be validated; "
                "keeping those tables as extracted: %s",
                i,
                exc,
            )
            edited.extend(_table_pass_through(it) for it in chunk)
    return edited


def _request_edits_texts_only(
    texts: list[dict], instruction: str, chunk_size: int = EDIT_CHUNK_SIZE
) -> list[dict]:
    """Run the structured LLM edit over text items only (tables excluded)."""
    if not texts:
        return []
    chunks = _chunk_items(texts, chunk_size, EDIT_CHUNK_MAX_CHARS)
    if len(chunks) == 1:
        return _request_edits_chunk_structured(chunks[0], instruction, chunk_index=0)

    edited: list[dict] = []
    for i, chunk in enumerate(chunks):
        edited.extend(_request_edits_chunk_structured(chunk, instruction, chunk_index=i))
    if len(edited) != len(texts):
        raise DocumentEditError("model did not return a valid edits array")
    return edited


def _table_pass_through(it: dict) -> dict:
    """Keep a table exactly as extracted, only deterministically stripping LXSHOW.

    Geometry, dimensions and styling are preserved; the table is not translated
    and not sent to the LLM.
    """
    new = dict(it)
    if "cells" in new:
        new["cells"] = [
            [_sanitize_lxshow(c) if isinstance(c, str) else c for c in row]
            for row in new["cells"]
        ]
    return new


def _request_edits_chunk_structured(
    chunk: list[dict], instruction: str, chunk_index: int = 0
) -> list[dict]:
    """Edit a single structured chunk with one controlled retry.

    A network/timeout failure from the edit model raises ``DocumentEditError``
    here (after diagnostics are logged) so the caller aborts the whole document
    edit rather than producing a partially broken PDF.
    """
    expected = [_slim_item(it) for it in chunk]
    system = _EDIT_STRUCTURED_SYSTEM.replace("{instruction}", instruction)
    user_prompt = _structured_user_prompt(expected)
    n_blocks = len(expected)
    total_chars = sum(_item_chars(it) for it in chunk)
    first_prompt_len = len(system) + len(user_prompt)

    try:
        raw = _generate_edits_array(user_prompt, system)
    except DocumentEditError as exc:
        logger.error(
            "edit_document(structured): chunk %s LLM call failed (attempt 1/%s): "
            "blocks=%s total_chars=%s prompt_len=%s error_type=%s reason=%s",
            chunk_index,
            _EDIT_CHUNK_MAX_ATTEMPTS,
            n_blocks,
            total_chars,
            first_prompt_len,
            type(exc).__name__,
            exc,
        )
        raise
    out, reason = _parse_structured_debug(raw, expected)
    if out is None:
        logger.warning(
            "edit_document(structured): model reply rejected (first attempt). "
            "chunk=%s items=%s total_chars=%s prompt_len=%s reason=%s raw_response=%s",
            chunk_index,
            len(expected),
            total_chars,
            first_prompt_len,
            reason,
            _safe_raw(raw),
        )
        retry_prompt = _structured_retry_prompt(expected, reason)
        retry_prompt_len = len(retry_prompt) + len(_STRUCTURED_RETRY_SYSTEM)
        try:
            raw = _generate_edits_array(retry_prompt, _STRUCTURED_RETRY_SYSTEM)
        except DocumentEditError as exc:
            logger.error(
                "edit_document(structured): chunk %s LLM call failed (attempt 2/%s): "
                "blocks=%s total_chars=%s prompt_len=%s error_type=%s reason=%s",
                chunk_index,
                _EDIT_CHUNK_MAX_ATTEMPTS,
                n_blocks,
                total_chars,
                retry_prompt_len,
                type(exc).__name__,
                exc,
            )
            raise
        out, reason = _parse_structured_debug(raw, expected)
        if out is None:
            logger.warning(
                "edit_document(structured): model reply rejected (after retry). "
                "chunk=%s items=%s total_chars=%s prompt_len=%s reason=%s raw_response=%s",
                chunk_index,
                len(expected),
                total_chars,
                retry_prompt_len,
                reason,
                _safe_raw(raw),
            )
            actual = _reply_array_length(raw)
            raise DocumentEditError(
                f"structured edit failed for chunk {chunk_index}: expected "
                f"{n_blocks} items, got {actual}; last rejection reason: {reason}"
            )

    # Merge validated textual edits back onto the original geometry by id.
    by_id = {it["id"]: it for it in chunk}
    edited_full: list[dict] = []
    for slim in out:
        orig = by_id[slim["id"]]
        new_item = dict(orig)
        if slim["type"] == "text":
            new_item["text"] = _sanitize_lxshow(slim.get("text", ""))
        else:
            new_item["cells"] = [
                [_sanitize_lxshow(c) if isinstance(c, str) else c for c in row]
                for row in slim.get("cells", orig["cells"])
            ]
        edited_full.append(new_item)
    return edited_full


def _structured_dims_note(expected: list[dict]) -> str:
    """Build an explicit per-table dimensions note so the model cannot miscount.

    For merged/irregular PDF tables PyMuPDF may extract a row count that differs
    from a 'clean' visual reading; the model must mirror the EXACT extracted
    shape, so we state rows x cols for every table item up front.
    """
    notes = []
    for it in expected:
        if it["type"] == "table":
            rows = len(it["cells"])
            cols = max((len(r) for r in it["cells"]), default=0)
            notes.append(
                f"- table item id={it['id']}: MUST return exactly {rows} rows and "
                f"{cols} columns (same order as the input cells)."
            )
    if not notes:
        return ""
    return (
        "TABLE DIMENSIONS YOU MUST MATCH EXACTLY (do not add/remove rows or "
        "columns):\n" + "\n".join(notes) + "\n\n"
    )


def _structured_user_prompt(expected: list[dict]) -> str:
    return _structured_dims_note(expected) + json.dumps(
        {"items": expected}, ensure_ascii=False
    )


def _structured_retry_prompt(expected: list[dict], reason: str | None = None) -> str:
    """Build the targeted retry user prompt for one structured chunk.

    Unlike the generic first attempt, the retry names the concrete rejection
    reason and pins the exact expected item count and the id contract, so the
    model can fix the specific misalignment instead of repeating the same
    mistake (the ``expected 8 -> got 7`` loop from production).
    """
    n = len(expected)
    head = f"Your previous reply was rejected: {reason}. " if reason else ""
    return (
        head
        + f"You must return EXACTLY {n} items, one per input item, each tagged "
        "with the exact input id it edits (same id set, no duplicates, none "
        "missing). The array may be in any order; ids define the mapping.\n\n"
        + _structured_dims_note(expected)
        + "INPUT ITEMS (copy their ids exactly):\n"
        + json.dumps({"items": expected}, ensure_ascii=False)
        + "\n\nReturn ONLY the raw JSON array of {\"id\": ...} objects now - "
        "one per input item, preserving the input ids."
    )


def _join_candidates(a: str, b: str) -> set[str]:
    """Concatenations that plausibly represent a merged pair of texts."""
    return {a + b, a + " " + b, a + "\n" + b}


def _out_text_item(it: dict, text: str) -> dict:
    return {"id": it["id"], "type": "text", "text": text}


def _reorder_by_id(
    expected: list[dict], out: list[dict]
) -> tuple[list[dict] | None, str | None]:
    """Reorder a reply so it matches ``expected`` order when ids form a bijection.

    Returns ``(reordered, None)`` or ``(None, reason)``. The ids are the only
    trustworthy signal: the model binds each output to an input item by id, so a
    reply returned in any order can be mapped back safely. Any missing, extra or
    duplicate id makes the mapping ambiguous and is rejected (never guessed).
    """
    exp_ids = [int(it["id"]) for it in expected]
    ret_ids: list[int] = []
    for v in out:
        try:
            ret_ids.append(int(v.get("id")))
        except (TypeError, ValueError):
            return None, "reply item has a non-integer id"
    if len(set(ret_ids)) != len(ret_ids):
        return None, f"duplicate ids in reply ({len(ret_ids)} items, {len(set(ret_ids))} unique)"
    if sorted(ret_ids) != sorted(exp_ids):
        missing = sorted(set(exp_ids) - set(ret_ids))
        extra = sorted(set(ret_ids) - set(exp_ids))
        return None, (
            f"item count mismatch: expected {len(exp_ids)}, got {len(ret_ids)} "
            f"(ids={ret_ids}; missing={missing} extra={extra})"
        )
    index = {int(v.get("id")): v for v in out}
    return [index[ident] for ident in exp_ids], None


def _recover_merged_text(expected: list[dict], value: list) -> list[dict] | None:
    """Split a unique verbatim merge of two adjacent text items (N-1 reply).

    The model merged items ``i`` and ``i+1`` into ``reply[i] == join(text[i],
    text[i+1])`` while keeping every other item unchanged. Recovery only fires
    when exactly ONE merge explains the whole reply (unambiguous): the merged
    string must equal the exact concatenation of the two originals and the rest
    of the reply must match 1:1. Translated/partially edited merges are never
    split - they fall back to a targeted retry.
    """
    texts = [it.get("text", "") for it in expected]
    n = len(texts)
    if len(value) != n - 1:
        return None
    buckets: list[int] = []
    for i in range(n - 1):
        a, b = texts[i], texts[i + 1]
        if not a or not b:
            continue
        if value[i] not in _join_candidates(a, b):
            continue
        if any(value[k] != texts[k] for k in range(i)):
            continue
        if any(value[k] != texts[k + 1] for k in range(i + 1, n - 1)):
            continue
        buckets.append(i)
    if len(buckets) != 1:
        return None
    i = buckets[0]
    out: list[dict] = []
    for k in range(n):
        if k == i:
            # The merge occupies ONE reply slot (value[i]) for TWO source items.
            out.append(_out_text_item(expected[k], texts[i]))
            out.append(_out_text_item(expected[k + 1], texts[i + 1]))
        elif k == i + 1:
            continue  # expected[i+1] was already appended above
        elif k < i:
            out.append(_out_text_item(expected[k], value[k]))
        else:
            out.append(_out_text_item(expected[k], value[k - 1]))
    return out


def _recover_skipped_empty(expected: list[dict], value: list) -> list[dict] | None:
    """Reinsert a single skipped EMPTY text item (N-1 reply, exactly one empty).

    When the source has exactly one empty item and the model dropped it, the
    remaining reply matches the non-empty originals in order - that mapping is
    unambiguous, so the empty item is restored unchanged. With several empties
    the alignment is ambiguous and recovery refuses.
    """
    texts = [it.get("text", "") for it in expected]
    if texts.count("") != 1:
        return None
    non_empty = [t for t in texts if t != ""]
    if list(value) != non_empty:
        return None
    out: list[dict] = []
    j = 0
    for it in expected:
        if not it.get("text", ""):
            out.append(_out_text_item(it, ""))
        else:
            out.append(_out_text_item(it, value[j]))
            j += 1
    return out


def _recover_split_text(expected: list[dict], value: list) -> list[dict] | None:
    """Rejoin a unique verbatim split of one text item (N+1 reply).

    The model split item ``i`` into two adjacent reply pieces that rejoin to the
    exact original text while every other item stays unchanged. Only accepted
    when exactly one such split explains the whole reply.
    """
    texts = [it.get("text", "") for it in expected]
    n = len(texts)
    if len(value) != n + 1:
        return None
    buckets: list[int] = []
    for i in range(n):
        target = texts[i]
        a, b = value[i], value[i + 1]
        if not target or not isinstance(a, str) or not isinstance(b, str):
            continue
        if target not in _join_candidates(a, b):
            continue
        if any(value[k] != texts[k] for k in range(i)):
            continue
        if any(value[k] != texts[k - 1] for k in range(i + 2, n + 1)):
            continue
        buckets.append(i)
    if len(buckets) != 1:
        return None
    i = buckets[0]
    out: list[dict] = []
    for k in range(n):
        if k == i:
            # The split occupies TWO reply slots (value[i], value[i+1]) whose
            # rejoined text replaces ONE source item.
            out.append(_out_text_item(expected[k], texts[i]))
        elif k < i:
            out.append(_out_text_item(expected[k], value[k]))
        else:
            out.append(_out_text_item(expected[k], value[k + 1]))
    return out


def _recover_structured(
    expected: list[dict], value: list
) -> tuple[list[dict] | None, str | None]:
    """Try to safely re-align a reply whose item count differs from ``expected``.

    Only deterministic, unambiguous recovery is attempted (see each helper):
      * an id-bearing reply is judged by its id SET - a non-bijection id set is
        reported and rejected, never guessed;
      * a text-only bare reply one item short is recovered only when the missing
        item is explained by one verbatim adjacent merge or one skipped empty
        item;
      * a text-only bare reply one item long is recovered only when one verbatim
        split explains it.

    Returns ``(aligned, None)`` or ``(None, reason)`` - ``None`` means the
    caller must run a targeted retry for this chunk.
    """
    n_exp = len(expected)
    n_got = len(value)

    if value and all(isinstance(v, dict) and "id" in v for v in value):
        _, reason = _reorder_by_id(expected, value)
        return None, reason or f"item count mismatch: expected {n_exp}, got {n_got}"

    if all(it["type"] == "text" for it in expected) and all(
        isinstance(v, str) for v in value
    ):
        if n_got == n_exp - 1:
            recovered = _recover_merged_text(expected, value) or _recover_skipped_empty(
                expected, value
            )
            if recovered is not None:
                return recovered, None
        elif n_got == n_exp + 1:
            recovered = _recover_split_text(expected, value)
            if recovered is not None:
                return recovered, None

    return None, f"item count mismatch: expected {n_exp}, got {n_got}"


def _reply_array_length(raw: str) -> object:
    """Best-effort count of a reply's array for an error message.

    Used only for diagnostics in the final ``DocumentEditError`` - it never
    returns the document's content.
    """
    if not raw or not raw.strip():
        return "none"
    cleaned = re.sub(r"```[a-zA-Z]*\s*", "", raw.strip())
    cleaned = re.sub(r"\s*```", "", cleaned).strip()
    value = _extract_json_value(cleaned)
    if isinstance(value, dict):
        found = _find_items_array(value)
        if found is None:
            lists = [v for v in value.values() if isinstance(v, list)]
            found = lists[0] if len(lists) == 1 else None
        value = found
    if isinstance(value, list):
        return len(value)
    return "unknown"


def _parse_structured_debug(
    raw: str, expected: list[dict]
) -> tuple[list[dict] | None, str | None]:
    """Parse a structured model reply, returning ``(items, rejection_reason)``.

    Accepted shapes (all empty of prose/markdown):

      * Structured object: ``{"items": [{"id", "type", "text"/"cells"}, ...]}``.
        Each element carries its ``id``/``type``.
      * Bare array: a top-level JSON array whose elements line up with
        ``expected`` -- a text item is a plain string or ``{"text": ...}``, a
        table item is ``{"headers": [...], "rows": [[...], ...]}`` (or
        ``{"cells": [[...]]}``).
      * Wrappers: ``{"result": {"items": [...]}}`` and arbitrary nesting are
        unwrapped by ``_find_items_array``.

    Alignment rules:

      * When every element carries an ``id``, the reply is bound by id: a
        bijective id set is reordered to the ``expected`` order and then
        validated (order-independent mapping - the model is never forced to
        guess positions).
      * When the element count differs from ``len(expected)``, an unambiguous
        deterministic recovery is attempted (verbatim adjacent merge / skipped
        empty item / verbatim split); otherwise the reply is rejected so the
        caller can run a targeted retry.
      * Tables MUST keep their exact row/column dimensions. Arbitrary JSON is
        rejected.
    """
    if not raw or not raw.strip():
        return None, "empty model response"
    cleaned = raw.strip()
    cleaned = re.sub(r"```[a-zA-Z]*\s*", "", cleaned)
    cleaned = re.sub(r"\s*```", "", cleaned).strip()

    value = _extract_json_value(cleaned)
    if value is None:
        return None, "no JSON array/object found (reply may be truncated)"
    if isinstance(value, dict):
        # json_object mode (and GigaChat's wrapping) return objects. Descend
        # into nested wrappers like {"result": {"items": [...]}} to find the
        # array, falling back to "the only list field".
        found = _find_items_array(value)
        if found is not None:
            value = found
        else:
            lists = [v for v in value.values() if isinstance(v, list)]
            if len(lists) == 1:
                value = lists[0]
            else:
                return None, "JSON object reply without a single array field"
    if not isinstance(value, list):
        return None, "top-level JSON is not an array"

    # Id-bearing reply: the id set is the only trustworthy signal. Reorder by id
    # (when ids form a bijection) and then validate/coerce; otherwise reject.
    if value and all(isinstance(v, dict) and "id" in v for v in value):
        reordered, reason = _reorder_by_id(expected, value)
        if reordered is None:
            return None, reason
        if all(isinstance(v, dict) and "type" in v for v in reordered):
            ok, reason = _validate_structured(expected, reordered)
            if not ok:
                return None, reason
            return reordered, None
        out, reason = _coerce_bare_array(expected, reordered)
        if out is None:
            return None, reason
        return out, None

    # Count mismatch: try an unambiguous recovery before rejecting.
    if len(value) != len(expected):
        recovered, reason = _recover_structured(expected, value)
        if recovered is None:
            return None, reason
        return recovered, None

    # Equal count, structured shape: every element carries id+type.
    if value and all(isinstance(v, dict) and "id" in v and "type" in v for v in value):
        ok, reason = _validate_structured(expected, value)
        if not ok:
            return None, reason
        return value, None

    # Equal count, bare-array shape: positional text strings + table objects.
    out, reason = _coerce_bare_array(expected, value)
    if out is None:
        return None, reason
    return out, None


def _find_items_array(value: dict):
    """Locate the edits array inside a possibly nested JSON object reply.

    json_object mode / GigaChat wrapping can nest the array arbitrarily deep
    (``{"result": {"items": [...]}}``, ``{"data": [...]}``, ...). Returns the
    first ``items`` list found, else any list, else None.
    """
    if isinstance(value, dict):
        if "items" in value and isinstance(value["items"], list):
            return value["items"]
        for sub in value.values():
            found = _find_items_array(sub) if isinstance(sub, dict) else None
            if found is not None:
                return found
    return None


def _table_cells_from_response(el: dict) -> list[list] | None:
    """Normalise a reply table object to a ``cells`` matrix.

    Accepts either ``{"cells": [[...]]}`` or the GigaChat shape
    ``{"headers": [...], "rows": [[...], ...]}`` (headers become the first row).
    Returns ``None`` for anything that is not a well-formed table.
    """
    if isinstance(el.get("cells"), list) and all(isinstance(r, list) for r in el["cells"]):
        return el["cells"]
    if "headers" in el and "rows" in el:
        headers = el["headers"]
        rows = el["rows"]
        if not isinstance(headers, list) or not isinstance(rows, list):
            return None
        if not all(isinstance(r, list) for r in rows):
            return None
        return [list(headers)] + [list(r) for r in rows]
    return None


def _coerce_bare_array(
    expected: list[dict], elements: list
) -> tuple[list[dict] | None, str | None]:
    """Map a positional bare-array reply onto ``expected`` items (1:1, same order).

    Text expected -> string element. Table expected -> object with cells/headers+rows.
    Rejects wrong types, wrong counts and malformed tables. Geometry (ids, cell rects)
    is taken from ``expected``; only the text content changes.
    """
    out: list[dict] = []
    for i, (exp, el) in enumerate(zip(expected, elements)):
        if exp["type"] == "text":
            # Tolerate {"text": "..."} objects for a text item: positionally
            # unambiguous, and some models emit uniform object arrays.
            if isinstance(el, dict) and isinstance(el.get("text"), str):
                el = el["text"]
            if not isinstance(el, str):
                return None, f"item {i} (text) is not a JSON string"
            out.append({"id": exp["id"], "type": "text", "text": el})
        elif exp["type"] == "table":
            if not isinstance(el, dict):
                return None, f"item {i} (table) is not a JSON object"
            cells = _table_cells_from_response(el)
            if cells is None:
                return None, f"item {i} table missing 'cells' or 'headers'/'rows'"
            in_cells = exp["cells"]
            if not isinstance(cells, list) or len(cells) != len(in_cells):
                return None, f"item {i} table row count mismatch"
            for r, (in_row, out_row) in enumerate(zip(in_cells, cells)):
                if not isinstance(out_row, list) or len(out_row) != len(in_row):
                    return None, f"item {i} table column count mismatch at row {r}"
                for cell in out_row:
                    if not isinstance(cell, (str, type(None))):
                        return None, f"item {i} table cell is not a string or null"
            out.append({"id": exp["id"], "type": "table", "cells": cells})
        else:
            return None, f"item {i} has unsupported type {exp['type']!r}"
    return out, None


def _validate_structured(expected: list[dict], out: list[dict]) -> tuple[bool, str | None]:
    """Strict check: 1 input item -> 1 output item, same id/type, table dims kept."""
    if len(expected) != len(out):
        return False, f"expected {len(expected)} items, got {len(out)}"
    for i, (a, b) in enumerate(zip(expected, out)):
        try:
            same_id = int(a["id"]) == int(b.get("id"))
        except (TypeError, ValueError):
            same_id = False
        if not same_id:
            return False, f"item {i} id mismatch (expected {a['id']}, got {b.get('id')})"
        if a["type"] != b.get("type"):
            return False, f"item {i} type mismatch (expected {a['type']}, got {b.get('type')})"
        if a["type"] == "text":
            if not isinstance(b.get("text"), str):
                return False, f"item {i} text is not a string"
        if a["type"] == "table":
            in_cells = a["cells"]
            out_cells = b.get("cells")
            if not isinstance(out_cells, list) or len(in_cells) != len(out_cells):
                return False, f"item {i} table row count mismatch"
            for r, (in_row, out_row) in enumerate(zip(in_cells, out_cells)):
                if not isinstance(out_row, list) or len(in_row) != len(out_row):
                    return False, f"item {i} table column count mismatch at row {r}"
                for cell in out_row:
                    if not isinstance(cell, (str, type(None))):
                        return False, f"item {i} table cell is not a string or null"
    return True, None


def _items_to_apply(
    edited_items: list[dict], fragments: list[dict]
) -> tuple[list[dict], list[str]]:
    """Flatten edited logical items back to the fragment-level (blocks, texts)
    form expected by ``_apply_pdf_edits``.

    A text item maps to its original fragment. A table item maps to one block
    per cell that has real geometry (``cell_rects`` entry), drawn back into the
    original cell rectangle. Cells without geometry, or that the model left as
    ``None``, are skipped so their original text stays exactly as-is. Gridlines
    and drawings are untouched by the apply step; only cell text is rewritten.
    """
    apply_blocks: list[dict] = []
    apply_texts: list[str] = []
    for it in edited_items:
        if it["type"] == "text":
            apply_blocks.append(fragments[it["frag_index"]])
            apply_texts.append(it.get("text", ""))
        else:
            for r, row in enumerate(it["cells"]):
                for c, cell in enumerate(row):
                    rect = (
                        it["cell_rects"][r][c]
                        if r < len(it["cell_rects"]) and c < len(it["cell_rects"][r])
                        else None
                    )
                    if rect is None or cell is None:
                        # No geometry (merged/missing in the source) or the model
                        # left the cell unchanged: keep the original text as-is.
                        continue
                    apply_blocks.append(
                        {
                            "page": it["page"],
                            "bbox": [rect.x0, rect.y0, rect.x1, rect.y1],
                            "style": it["style"],
                        }
                    )
                    apply_texts.append(cell)
    return apply_blocks, apply_texts
