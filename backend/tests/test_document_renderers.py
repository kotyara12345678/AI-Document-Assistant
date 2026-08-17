"""Renderers: real .docx / .odt / .pdf files are produced from a DocumentSpec
and parsed back to prove title, headings, paragraphs, lists, tables and
Cyrillic survive the round-trip."""

import io

from app.schemas.document_spec import DocumentSpec
from app.services.docx_renderer import render_docx
from app.services.odt_renderer import render_odt
from app.services.pdf_renderer import render_pdf


def _spec() -> DocumentSpec:
    return DocumentSpec.model_validate(
        {
            "title": "Трудовой договор № 12",
            "author": "ООО Ромашка",
            "blocks": [
                {"type": "heading", "level": 1, "text": "1. Общие положения"},
                {"type": "paragraph", "text": "Стороны заключили настоящий договор."},
                {"type": "heading", "level": 2, "text": "Предмет договора"},
                {"type": "paragraph", "text": "Работник выполняет обязанности менеджера."},
                {"type": "list", "items": ["Знание русского языка", "Опыт от 2 лет"]},
                {
                    "type": "table",
                    "headers": ["Показатель", "Значение"],
                    "rows": [["Оклад", "150 000 руб."], ["Срок", "1 год"]],
                },
            ],
        }
    )


def _docx_paragraph_texts(data: bytes) -> list[str]:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    return [p.text for p in doc.paragraphs]


def _odt_text(data: bytes) -> str:
    from odf import teletype
    from odf.opendocument import load

    doc = load(io.BytesIO(data))
    return " ".join(teletype.extractText(doc.text).split())


# ---------------------------------------------------------------- DOCX


def test_docx_is_a_valid_archive_and_contains_title():
    data = render_docx(_spec())
    assert data[:2] == b"PK"
    with io.BytesIO(data) as stream:
        import zipfile

        with zipfile.ZipFile(stream) as zf:
            assert "word/document.xml" in zf.namelist()
    texts = " | ".join(_docx_paragraph_texts(data))
    assert "Трудовой договор № 12" in texts


def test_docx_preserves_headings_and_paragraphs():
    data = render_docx(_spec())
    texts = _docx_paragraph_texts(data)
    joined = " | ".join(texts)
    for expected in (
        "1. Общие положения",
        "Предмет договора",
        "Стороны заключили настоящий договор.",
        "Работник выполняет обязанности менеджера.",
    ):
        assert expected in joined, f"missing in docx: {expected!r}"


def test_docx_preserves_lists_and_tables():
    data = render_docx(_spec())
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    list_texts = {p.text for p in doc.paragraphs if p.style.name == "List Bullet"}
    assert {"Знание русского языка", "Опыт от 2 лет"} <= list_texts

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert table.rows[0].cells[0].text == "Показатель"
    assert table.rows[1].cells[1].text == "150 000 руб."


def test_docx_numbered_list_uses_list_number_style():
    spec = DocumentSpec.model_validate(
        {
            "title": "Порядок действий",
            "blocks": [{"type": "list", "ordered": True, "items": ["Шаг 1", "Шаг 2"]}],
        }
    )
    data = render_docx(spec)
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    assert {p.text for p in doc.paragraphs if p.style.name == "List Number"} == {
        "Шаг 1",
        "Шаг 2",
    }


def test_docx_handles_empty_and_long_texts():
    spec = DocumentSpec.model_validate(
        {
            "title": "Пустой",
            "blocks": [
                {"type": "list", "items": []},
                {"type": "paragraph", "text": "длинный " * 200},
            ],
        }
    )
    data = render_docx(spec)
    assert data[:2] == b"PK"


# ---------------------------------------------------------------- ODT


def test_odt_is_a_valid_archive_and_contains_title():
    data = render_odt(_spec())
    assert data[:2] == b"PK"
    with io.BytesIO(data) as stream:
        import zipfile

        with zipfile.ZipFile(stream) as zf:
            assert "content.xml" in zf.namelist()
    text = _odt_text(data)
    assert "Трудовой договор № 12" in text


def test_odt_preserves_headings_and_paragraphs():
    data = render_odt(_spec())
    text = _odt_text(data)
    for expected in (
        "1. Общие положения",
        "Предмет договора",
        "Стороны заключили настоящий договор.",
        "Работник выполняет обязанности менеджера.",
    ):
        assert expected in text, f"missing in odt: {expected!r}"


def test_odt_preserves_lists_and_tables():
    data = render_odt(_spec())
    text = _odt_text(data)
    assert "Знание русского языка" in text
    assert "Опыт от 2 лет" in text
    assert "Оклад" in text
    assert "150 000 руб." in text


def test_odt_numbered_list_marks_items():
    spec = DocumentSpec.model_validate(
        {
            "title": "Порядок действий",
            "blocks": [{"type": "list", "ordered": True, "items": ["Шаг 1", "Шаг 2"]}],
        }
    )
    data = render_odt(spec)
    assert "Шаг 1" in _odt_text(data)
    assert "Шаг 2" in _odt_text(data)


def test_odt_handles_empty_and_long_texts():
    spec = DocumentSpec.model_validate(
        {
            "title": "Пустой",
            "blocks": [
                {"type": "list", "items": []},
                {"type": "paragraph", "text": "длинный " * 200},
            ],
        }
    )
    data = render_odt(spec)
    assert data[:2] == b"PK"
    assert "длинный" in _odt_text(data)


# ---------------------------------------------------------------- PDF


def _pdf_text(data: bytes) -> str:
    from app.services.extraction import extract_text

    return extract_text(data, "pdf")


def test_pdf_starts_with_magic_and_contains_title():
    data = render_pdf(_spec())
    assert data.startswith(b"%PDF")
    assert "Трудовой договор № 12" in _pdf_text(data)


def test_pdf_preserves_headings_paragraphs_lists_and_tables():
    data = render_pdf(_spec())
    text = _pdf_text(data)
    for expected in (
        "1. Общие положения",
        "Предмет договора",
        "Стороны заключили настоящий договор.",
        "Работник выполняет обязанности менеджера.",
        "Знание русского языка",
        "Опыт от 2 лет",
        "Оклад",
        "150 000 руб.",
    ):
        assert expected in text, f"missing in pdf: {expected!r}"


def test_pdf_handles_empty_and_long_texts():
    spec = DocumentSpec.model_validate(
        {
            "title": "Пустой",
            "blocks": [
                {"type": "list", "items": []},
                {"type": "paragraph", "text": "длинный " * 200},
            ],
        }
    )
    data = render_pdf(spec)
    assert data.startswith(b"%PDF")
    assert "длинный" in _pdf_text(data)


def test_pdf_passes_validation_and_rejects_missing_text():
    import pytest

    from app.services.document_validation import validate_document_bytes
    from app.services.errors import RendererError

    spec = _spec()
    data = render_pdf(spec)
    validate_document_bytes(data, "pdf", spec)  # must not raise

    other = DocumentSpec.model_validate(
        {
            "title": "Другой",
            "blocks": [{"type": "paragraph", "text": "этого абзаца точно нет"}],
        }
    )
    with pytest.raises(RendererError):
        validate_document_bytes(data, "pdf", other)


# ---------------------------------------------------------------- MD / TXT


def test_txt_is_valid_plain_utf8_and_preserves_text():
    from app.services.document_renderer import render_txt

    data = render_txt(_spec())
    text = data.decode("utf-8")
    assert text.startswith("Трудовой договор № 12")
    for expected in (
        "1. Общие положения",
        "Стороны заключили настоящий договор.",
        "- Знание русского языка",
        "Показатель | Значение",
        "Оклад | 150 000 руб.",
    ):
        assert expected in text, f"missing in txt: {expected!r}"


def test_md_is_valid_markdown_and_preserves_blocks():
    from app.services.document_renderer import render_md

    data = render_md(_spec())
    text = data.decode("utf-8")
    assert text.startswith("# Трудовой договор № 12")
    for expected in (
        "# 1. Общие положения",
        "Стороны заключили настоящий договор.",
        "- Знание русского языка",
        "| Показатель | Значение |",
        "| Оклад | 150 000 руб. |",
    ):
        assert expected in text, f"missing in md: {expected!r}"


def test_md_and_txt_pass_validation_and_reject_missing_text():
    import pytest

    from app.services.document_renderer import render_md, render_txt
    from app.services.document_validation import validate_document_bytes
    from app.services.errors import RendererError

    spec = _spec()
    for fmt, render in (("txt", render_txt), ("md", render_md)):
        data = render(spec)
        validate_document_bytes(data, fmt, spec)  # must not raise

    other = DocumentSpec.model_validate(
        {
            "title": "Другой",
            "blocks": [{"type": "paragraph", "text": "этого абзаца точно нет"}],
        }
    )
    for fmt, render in (("txt", render_txt), ("md", render_md)):
        with pytest.raises(RendererError):
            validate_document_bytes(render(spec), fmt, other)
