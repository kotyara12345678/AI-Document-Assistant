"""Tests for the minimal agent layer (POST /api/agent).

The GigaChat client is mocked at the ``chat_with_functions`` boundary, so the
agent loop, the tool execution against the real retrieval pipeline and the
message round-trip are all exercised in-process against real PostgreSQL and
Qdrant.
"""

import json
import uuid
from pathlib import Path

from fastapi.testclient import TestClient

from app.core.config import settings
from app.main import app
from app.services import gemini
from app.services.agent import agent_service

API_PREFIX = "/api"


def _upload(client: TestClient, filename: str, content: bytes):
    files = {"file": (filename, content)}
    return client.post(f"{API_PREFIX}/documents/upload", files=files)


def _get_document(document_id: int):
    from app.database.session import SessionLocal
    from app.models.document import Document

    db = SessionLocal()
    try:
        return db.query(Document).filter(Document.id == document_id).first()
    finally:
        db.close()


def _tool_call_message(name="search_documents", arguments='{"query": "про что документ"}'):
    return {
        "role": "assistant",
        "content": None,
        "function_call": {"name": name, "arguments": arguments},
    }


def _scripted_functions(monkeypatch, script):
    """Queue (message, state_id) pairs for successive chat_with_functions calls.

    Records every (messages, functions, functions_state_id) tuple so tests can
    assert what the agent actually sent to (and received from) the model.
    """
    calls = []

    def fake(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        calls.append(
            {
                "messages": messages,
                "functions": functions,
                "function_call": function_call,
                "functions_state_id": functions_state_id,
            }
        )
        return script.pop(0)

    monkeypatch.setattr(gemini, "chat_with_functions", fake)
    return calls


def test_agent_answers_without_tool_call(client, monkeypatch):
    """A plain query that needs no documents -> direct answer, no tools."""
    calls = _scripted_functions(
        monkeypatch, [({"content": "Привет! Чем помочь с документами?"}, None)]
    )

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "привет"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Привет! Чем помочь с документами?"
    assert data["tool_calls"] == []
    assert data["tool_results"] == []

    # The function spec must have been advertised even when unused.
    assert len(calls) == 1
    assert calls[0]["functions"] == agent_service.functions_spec()
    assert calls[0]["function_call"] == "auto"
    assert calls[0]["functions_state_id"] is None


def test_agent_calls_search_documents(client, monkeypatch):
    """When the model requests search_documents, the real pipeline runs and the
    compact result (id, filename, score, snippet) reaches the model."""
    marker = f"AGT{uuid.uuid4().hex[:6]}"
    text = (
        f"План продаж дронов {marker}. Рекламный бюджет 900000 рублей в квартал."
    ) * 20
    resp = _upload(client, "agent_doc.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()[0]["id"]

    tool_msg = _tool_call_message(arguments='{"query": "какой рекламный бюджет"}')
    final_msg = {"content": "Рекламный бюджет составляет 900000 рублей."}
    calls = _scripted_functions(
        monkeypatch, [(tool_msg, "state-1"), (final_msg, None)]
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": f"какой рекламный бюджет дронов {marker}"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Рекламный бюджет составляет 900000 рублей."

    # The tool call was recorded with its parsed arguments.
    assert len(data["tool_calls"]) == 1
    assert data["tool_calls"][0]["name"] == "search_documents"
    assert data["tool_calls"][0]["arguments"] == {"query": "какой рекламный бюджет"}

    # The tool result is a compact, structured JSON string.
    assert len(data["tool_results"]) == 1
    result = data["tool_results"][0]
    assert result["tool_call_id"] == "search_documents"
    assert result["name"] == "search_documents"
    assert "agent_doc.txt" in result["content"]
    assert str(document_id) in result["content"]
    assert '"score"' in result["content"]
    assert '"snippet"' in result["content"]
    assert marker in result["content"]

    # Exactly two LLM calls: one requesting the tool, one producing the answer.
    assert len(calls) == 2


def test_agent_threads_functions_state_id(client, monkeypatch):
    """GigaChat's functions_state_id from round one is echoed into round two."""
    tool_msg = _tool_call_message(arguments='{"query": "что там"}')
    final_msg = {"content": "Итог."}
    calls = _scripted_functions(
        monkeypatch, [(tool_msg, "state-abc"), (final_msg, None)]
    )

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "что там"})
    assert resp.status_code == 200, resp.text

    assert calls[1]["functions_state_id"] == "state-abc"


def test_agent_passes_tool_result_back_to_model(client, monkeypatch):
    """The assistant function_call message and the role:function result are both
    sent back to the model in the correct order."""
    marker = f"PAS{uuid.uuid4().hex[:6]}"
    text = (f"Документ про запуск {marker}. Дата релиза 1 сентября.") * 20
    _upload(client, "passback.txt", text.encode("utf-8"))

    tool_msg = _tool_call_message(arguments='{"query": "когда релиз"}')
    final_msg = {"content": "Релиз 1 сентября."}
    calls = _scripted_functions(
        monkeypatch, [(tool_msg, "state-2"), (final_msg, None)]
    )

    resp = client.post(f"{API_PREFIX}/agent", json={"question": f"когда релиз {marker}"})
    assert resp.status_code == 200, resp.text

    second = calls[1]["messages"]
    roles = [m["role"] for m in second]
    assert roles == ["system", "user", "assistant", "function"], roles

    assistant_msg = second[2]
    assert assistant_msg.get("function_call") == tool_msg["function_call"]

    tool_msg_out = second[3]
    assert tool_msg_out["role"] == "function"
    assert tool_msg_out["name"] == "search_documents"
    assert "passback.txt" in tool_msg_out["content"]
    assert marker in tool_msg_out["content"]


def test_agent_handles_search_error(client, monkeypatch):
    """A failing retrieval must not crash the loop: the model gets an error
    object in the tool result and still produces a final answer."""
    def boom(user_id, query, document_ids):
        raise RuntimeError("qdrant exploded")

    monkeypatch.setattr(agent_service, "_search_documents", boom)

    tool_msg = _tool_call_message(arguments='{"query": "что там в документах"}')
    final_msg = {"content": "Не удалось выполнить поиск по документам."}
    _scripted_functions(monkeypatch, [(tool_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "что в документах"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Не удалось выполнить поиск по документам."
    assert len(data["tool_results"]) == 1
    assert "error" in data["tool_results"][0]["content"]
    assert "qdrant exploded" in data["tool_results"][0]["content"]


def test_agent_unknown_tool_is_reported(client, monkeypatch):
    """An unexpected tool name is reported back instead of crashing."""
    unknown = _tool_call_message(name="write_file", arguments="{}")
    final_msg = {"content": "Такого инструмента нет."}
    _scripted_functions(monkeypatch, [(unknown, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "запиши файл"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert len(data["tool_results"]) == 1
    assert data["tool_results"][0]["name"] == "write_file"
    assert "Unknown tool" in data["tool_results"][0]["content"]


def test_agent_respects_user_permissions(client, monkeypatch, register_user):
    """A user can never search documents owned by another user: the tool runs
    scoped to the caller, so another user's document never appears."""
    marker = f"PRM{uuid.uuid4().hex[:6]}"
    text = (f"Секретные данные {marker}. Пароль от сервера = s3cr3t.") * 20
    resp = _upload(client, "private.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    private_id = resp.json()[0]["id"]

    tool_msg = _tool_call_message(arguments=f'{{"query": "{marker} секретный пароль"}}')
    final_msg = {"content": "Не нашёл такой информации."}
    _scripted_functions(monkeypatch, [(tool_msg, "s"), (final_msg, None)])

    with TestClient(app) as other:
        info = register_user(other)
        other.headers.update({"Authorization": f"Bearer {info['token']}"})
        resp = other.post(
            f"{API_PREFIX}/agent",
            json={"question": f"что за секрет {marker}"},
        )
        assert resp.status_code == 200, resp.text
        data = resp.json()

        assert len(data["tool_results"]) == 1
        content = data["tool_results"][0]["content"]
        assert "private.txt" not in content
        assert str(private_id) not in content
        assert marker not in content


def test_agent_requires_authentication():
    """No bearer token -> HTTP 401, same as every other protected endpoint."""
    with TestClient(app) as c:
        resp = c.post(f"{API_PREFIX}/agent", json={"question": "привет"})
    assert resp.status_code == 401, resp.text


def test_agent_degrades_when_gigachat_fails(client, monkeypatch):
    """An LLM outage degrades to an honest message instead of a 500."""
    def failing(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        raise gemini.GeminiError("boom")

    monkeypatch.setattr(gemini, "chat_with_functions", failing)

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "привет"})
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert "unavailable" in data["answer"].lower()


def _insert_document(user_id: int, filename: str, content: str, file_type: str = "txt") -> int:
    """Insert a Document row directly (no embedding) for focused tool tests."""
    from app.database.session import SessionLocal
    from app.models.document import Document

    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename="stored.txt",
            original_filename=filename,
            file_type=file_type,
            file_size=len(content.encode("utf-8")),
            filepath="/data/uploads/stored.txt",
            content=content,
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc.id
    finally:
        db.close()


def test_agent_lists_all_documents_not_just_chat_memory(client, monkeypatch, user_id):
    """«Перечисли мне список всех моих файлов» must trigger list_documents and
    the tool must return EVERY document of the user (all 15), because the DB
    listing — not search hits or chat memory — is the source of truth."""
    names = [f"file_{index:02d}.txt" for index in range(1, 16)]
    for name in names:
        _insert_document(user_id, name, f"document {name} contents")

    list_msg = _tool_call_message("list_documents", "{}")
    final_answer = "\n".join(f"- {name}" for name in names)
    final_msg = {"content": final_answer}
    calls = _scripted_functions(monkeypatch, [(list_msg, "s"), (final_msg, None)])

    resp = client.post(
        f"{API_PREFIX}/agent", json={"question": "Перечисли мне список всех моих файлов"}
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    # list_documents was called (not search_documents/RAG).
    assert len(data["tool_calls"]) == 1
    assert data["tool_calls"][0]["name"] == "list_documents"

    # The tool was advertised to the model and the listing rule is in the prompt.
    names_advertised = {f["name"] for f in calls[0]["functions"]}
    assert "list_documents" in names_advertised
    assert "list_documents" in calls[0]["messages"][0]["content"]

    # The tool result contains ALL 15 documents with id/filename/type/date.
    result = data["tool_results"][0]
    assert result["name"] == "list_documents"
    content = json.loads(result["content"])
    assert len(content) == 15
    for entry in content:
        assert "document_id" in entry
        assert "filename" in entry
        assert "type" in entry
        assert "created_at" in entry
    filenames = {entry["filename"] for entry in content}
    assert filenames == set(names)

    # The final answer the model produced enumerates every file.
    assert data["answer"] == final_answer


def test_agent_calls_read_document(client, monkeypatch, user_id):
    """When the model requests read_document, the stored document content,
    filename and type are returned to the model in a structured result."""
    marker = f"RED{uuid.uuid4().hex[:6]}"
    text = (f"Служебная записка {marker}. Должность Сергея Юрьевича — директор.") * 5
    document_id = _insert_document(user_id, "zapiska.txt", text)

    read_msg = _tool_call_message("read_document", {"document_id": document_id})
    final_msg = {"content": "Сергей Юрьевич — директор."}
    calls = _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": f"кто такой Сергей {marker}"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Сергей Юрьевич — директор."

    # The agent advertised both tools to the model.
    names = {f["name"] for f in calls[0]["functions"]}
    assert {"search_documents", "read_document"} <= names

    # The read result is structured and carries the document metadata + text.
    assert len(data["tool_calls"]) == 1
    assert data["tool_calls"][0]["name"] == "read_document"
    assert data["tool_calls"][0]["arguments"] == {"document_id": document_id}

    assert len(data["tool_results"]) == 1
    result = json.loads(data["tool_results"][0]["content"])
    assert result["document_id"] == document_id
    assert result["filename"] == "zapiska.txt"
    assert result["file_type"] == "txt"
    assert result["file_size"] == len(text.encode("utf-8"))
    assert result["content_length"] == len(text)
    assert result["truncated"] is False
    assert marker in result["text"]
    assert "директор" in result["text"]


def test_read_document_respects_user_permissions(client, monkeypatch, user_id, register_user):
    """A user can never read another user's document: the tool returns the
    same 'document not found' error as for a nonexistent id and leaks nothing."""
    marker = f"SEC{uuid.uuid4().hex[:6]}"
    text = (f"Секретный устав {marker}. Пароль = s3cr3t.") * 5
    private_id = _insert_document(user_id, "private.txt", text)

    read_msg = _tool_call_message("read_document", {"document_id": private_id})
    final_msg = {"content": "Не могу прочитать документ."}
    _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])

    with TestClient(app) as other:
        info = register_user(other)
        other.headers.update({"Authorization": f"Bearer {info['token']}"})
        resp = other.post(
            f"{API_PREFIX}/agent",
            json={"question": "что за документ"},
        )
        assert resp.status_code == 200, resp.text
        data = resp.json()

        assert len(data["tool_results"]) == 1
        result = json.loads(data["tool_results"][0]["content"])
        assert result["error"] == "document not found"
        assert "private.txt" not in result
        assert marker not in json.dumps(result, ensure_ascii=False)


def test_read_document_unknown_document(client, monkeypatch):
    """A nonexistent document_id yields a clean error result, not a crash."""
    read_msg = _tool_call_message("read_document", {"document_id": 999_999})
    final_msg = {"content": "Такого документа нет."}
    _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "прочитай 999999"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert len(data["tool_results"]) == 1
    result = json.loads(data["tool_results"][0]["content"])
    assert result["error"] == "document not found"
    assert result["document_id"] == 999_999


def test_read_document_bad_arguments(client, monkeypatch):
    """A missing/non-numeric document_id is reported instead of crashing."""
    for bad in ({}, {"document_id": "не число"}, {"document_id": None}):
        read_msg = _tool_call_message("read_document", bad)
        final_msg = {"content": "Некорректные аргументы."}
        _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])

        resp = client.post(f"{API_PREFIX}/agent", json={"question": "прочитай"})
        assert resp.status_code == 200, resp.text
        result = json.loads(resp.json()["tool_results"][0]["content"])
        assert "document_id" in result["error"]


def test_agent_search_then_read_then_answer(client, monkeypatch, user_id):
    """The full tool sequence the user cares about:
    search_documents -> read_document -> final answer."""
    marker = f"SEA{uuid.uuid4().hex[:6]}"
    text = (f"Кадровые данные {marker}. Должность Сергея Юрьевича — заместитель директора.") * 20
    resp = _upload(client, "kadry.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()[0]["id"]

    search_msg = _tool_call_message(
        "search_documents", {"query": f"Сергей Юрьевич должность {marker}"}
    )
    read_msg = _tool_call_message("read_document", {"document_id": document_id})
    final_msg = {"content": "Сергей Юрьевич — заместитель директора."}
    calls = _scripted_functions(
        monkeypatch, [(search_msg, "s1"), (read_msg, "s2"), (final_msg, None)]
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": f"Найди информацию о Сергее Юрьевиче и расскажи его должность {marker}"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Сергей Юрьевич — заместитель директора."

    # Both tool calls are recorded in order.
    names = [call["name"] for call in data["tool_calls"]]
    assert names == ["search_documents", "read_document"]
    assert data["tool_calls"][1]["arguments"] == {"document_id": document_id}

    # Exactly three LLM calls: search, read, final answer.
    assert len(calls) == 3

    # The second round receives the functions_state_id threaded from the first.
    assert calls[1]["functions_state_id"] == "s1"

    # Round three sees the accumulated conversation: user, assistant(2), function(2).
    third_roles = [m["role"] for m in calls[2]["messages"]]
    assert third_roles == [
        "system", "user", "assistant", "function", "assistant", "function",
    ], third_roles


def test_read_document_error_is_reported(client, monkeypatch):
    """A failing read must not crash the loop: the model gets an error object
    and still produces a final answer."""
    def boom(arguments, user_id):
        raise RuntimeError("db exploded")

    monkeypatch.setattr(agent_service, "_read_document", boom)

    read_msg = _tool_call_message("read_document", {"document_id": 1})
    final_msg = {"content": "Не удалось прочитать документ."}
    _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "прочитай документ"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Не удалось прочитать документ."
    assert len(data["tool_results"]) == 1
    assert "error" in data["tool_results"][0]["content"]
    assert "db exploded" in data["tool_results"][0]["content"]


def test_read_document_limits_context(client, monkeypatch, user_id):
    """Oversized documents are returned in bounded windows: the first call
    truncates to AGENT_READ_MAX_CHARS and flags it, a follow-up call with an
    offset reads the remaining portion."""
    max_chars = settings.AGENT_READ_MAX_CHARS
    text = ("Параграф о полётах дронов и квотах. " * 300) + "КОНЕЦ ДОКУМЕНТА"
    assert max_chars < len(text) < 2 * max_chars
    document_id = _insert_document(user_id, "big.txt", text)

    # First read: no offset -> first window, truncated.
    read1 = _tool_call_message("read_document", {"document_id": document_id})
    final1 = {"content": "первая часть ок"}
    _scripted_functions(monkeypatch, [(read1, "s"), (final1, None)])
    resp = client.post(f"{API_PREFIX}/agent", json={"question": "первая часть"})
    assert resp.status_code == 200, resp.text
    first = json.loads(resp.json()["tool_results"][0]["content"])
    assert first["truncated"] is True
    assert first["length"] == max_chars
    assert first["text"] == text[:max_chars]
    assert first["offset"] == 0

    # Second read: offset = max_chars -> remainder, not truncated.
    read2 = _tool_call_message(
        "read_document", {"document_id": document_id, "offset": max_chars}
    )
    final2 = {"content": "вторая часть ок"}
    _scripted_functions(monkeypatch, [(read2, "s"), (final2, None)])
    resp = client.post(f"{API_PREFIX}/agent", json={"question": "вторая часть"})
    assert resp.status_code == 200, resp.text
    second = json.loads(resp.json()["tool_results"][0]["content"])
    assert second["truncated"] is False
    assert second["offset"] == max_chars
    assert second["text"] == text[max_chars:]


# ---------------------------------------------------------------- create_document


def _create_call(document_spec: dict, output_format: str = "docx"):
    return _tool_call_message(
        "create_document",
        {"document_spec": document_spec, "output_format": output_format},
    )


def test_agent_calls_create_document(client, monkeypatch, user_id):
    """create_document produces a real, parseable docx owned by the caller."""
    create_msg = _create_call(
        {
            "title": "Трудовой договор",
            "author": "ООО Ромашка",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Общие положения"},
                {"type": "paragraph", "text": "Стороны заключили договор."},
            ],
        }
    )
    final_msg = {"content": "Договор создан и сохранён."}
    calls = _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай договор"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Договор создан и сохранён."
    names = {f["name"] for f in calls[0]["functions"]}
    assert "create_document" in names

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    assert result["file_type"] == "docx"
    assert result["filename"].endswith(".docx")

    # CONFIRMATION QUALITY: the tool result carries factual confirmation data
    # (real file name, format, download url) the model must cite verbatim.
    assert result["download_url"].startswith(f"{API_PREFIX}/documents/")
    assert result["confirmation"]
    assert "Документ" in result["confirmation"]
    assert result["filename"] in result["confirmation"]
    assert "DOCX" in result["confirmation"]

    doc = _get_document(result["document_id"])
    assert doc is not None
    assert doc.user_id == user_id
    assert doc.file_type == "docx"
    assert result["file_size"] == doc.file_size
    assert result["document_id"] == doc.id

    # A real file on disk that python-docx can reopen, with Cyrillic intact.
    path = Path(doc.filepath)
    assert path.is_file()
    assert path.read_bytes()[:2] == b"PK"
    from docx import Document as DocxDocument

    parsed = DocxDocument(str(path))
    joined = " | ".join(p.text for p in parsed.paragraphs)
    assert "Трудовой договор" in joined
    assert "Стороны заключили договор." in joined
    assert "Общие положения" in joined

    # The content column stores the plain-text form (searchable/readable).
    assert "Трудовой договор" in doc.content
    assert "Стороны заключили договор." in doc.content


def test_agent_calls_create_document_odt(client, monkeypatch, user_id):
    """The same spec renders as a real, parseable odt."""
    create_msg = _create_call(
        {
            "title": "Отчёт по проекту",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Итоги"},
                {"type": "paragraph", "text": "Проект завершён успешно."},
            ],
        },
        output_format="odt",
    )
    final_msg = {"content": "ODT готов."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай odt"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is True
    assert result["file_type"] == "odt"
    assert result["filename"].endswith(".odt")

    doc = _get_document(result["document_id"])
    assert doc.user_id == user_id
    path = Path(doc.filepath)
    assert path.is_file()
    assert path.read_bytes()[:2] == b"PK"

    from odf import teletype
    from odf.opendocument import load

    loaded = load(str(path))
    text = " ".join(teletype.extractText(loaded.text).split())
    assert "Отчёт по проекту" in text
    assert "Проект завершён успешно." in text


def test_create_document_validates_spec(client, monkeypatch):
    """An invalid spec (missing title) yields a safe structured error and no file."""
    create_msg = _create_call({"blocks": []})
    final_msg = {"content": "Создать не удалось."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is False
    assert "title" in result["error"].lower()


def test_create_document_rejects_oversized_spec(client, monkeypatch):
    """An oversized spec is rejected before any file is generated."""
    create_msg = _create_call(
        {
            "title": "Док",
            "blocks": [{"type": "paragraph", "text": "x" * (settings.AGENT_DOCUMENT_MAX_LINE_CHARS + 1)}],
        }
    )
    final_msg = {"content": "Слишком большой."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай большой"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is False
    assert "too long" in result["error"]


def test_create_document_bad_format(client, monkeypatch):
    """An unsupported output format is rejected."""
    create_msg = _create_call({"title": "Док"}, output_format="pptx")
    final_msg = {"content": "Не получилось."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай pptx"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is False
    assert "unsupported output format" in result["error"]


def test_create_document_ignores_user_id_from_arguments(client, monkeypatch, user_id):
    """user_id is never trusted from tool arguments: ownership comes from the
    authenticated request context."""
    create_msg = _tool_call_message(
        "create_document",
        {
            "document_spec": {"title": "Секретный документ", "blocks": [{"type": "paragraph", "text": "демо"}]},
            "output_format": "docx",
            "user_id": 999_999,
        },
    )
    final_msg = {"content": "готово"}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is True

    doc = _get_document(result["document_id"])
    assert doc is not None
    assert doc.user_id == user_id


def test_create_document_ownership(client, monkeypatch, register_user):
    """A document created by user B is never readable by user A."""
    with TestClient(app) as other:
        info = register_user(other)
        other.headers.update({"Authorization": f"Bearer {info['token']}"})
        create_msg = _create_call({"title": "Приватный договор", "blocks": [{"type": "paragraph", "text": "демо"}]})
        final_msg = {"content": "ок"}
        _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

        resp = other.post(f"{API_PREFIX}/agent", json={"question": "создай"})
        assert resp.status_code == 200, resp.text
        created_id = json.loads(resp.json()["tool_results"][0]["content"])["document_id"]

    # User A's read_document against B's created document -> not found.
    read_msg = _tool_call_message("read_document", {"document_id": created_id})
    final_msg = {"content": "Нет доступа."}
    _scripted_functions(monkeypatch, [(read_msg, "s"), (final_msg, None)])
    resp = client.post(f"{API_PREFIX}/agent", json={"question": "прочитай"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["error"] == "document not found"


def test_agent_search_read_create_answer(client, monkeypatch, user_id):
    """The full user scenario: search -> read -> create -> final answer."""
    marker = f"GEN{uuid.uuid4().hex[:6]}"
    text = (f"Данные сотрудника {marker}. Сергей Юрьевич, должность — директор.") * 20
    resp = _upload(client, "data.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()[0]["id"]

    search_msg = _tool_call_message(
        "search_documents", {"query": f"Сергей Юрьевич должность {marker}"}
    )
    read_msg = _tool_call_message("read_document", {"document_id": document_id})
    create_msg = _create_call(
        {
            "title": "Договор",
            "blocks": [{"type": "paragraph", "text": f"Сергей Юрьевич — директор. {marker}"}],
        }
    )
    final_msg = {"content": "Договор создан на основе найденных данных."}
    calls = _scripted_functions(
        monkeypatch,
        [(search_msg, "s1"), (read_msg, "s2"), (create_msg, "s3"), (final_msg, None)],
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": f"Создай документ на основе найденных данных {marker}"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Договор создан на основе найденных данных."
    names = [call["name"] for call in data["tool_calls"]]
    assert names == ["search_documents", "read_document", "create_document"]

    create_result = json.loads(data["tool_results"][-1]["content"])
    assert create_result["success"] is True
    assert create_result["file_type"] == "docx"

    created = _get_document(create_result["document_id"])
    assert created is not None
    assert created.user_id == user_id
    assert Path(created.filepath).is_file()
    assert marker in created.content

    # Three tool rounds consumed, the final answer comes from the last call.
    assert len(calls) == 4


# ------------------------------------------------------- create policy


def test_agent_creates_any_document_default_docx(client, monkeypatch, user_id):
    """'сгенерируй любой документ' -> the agent calls create_document with a
    minimal demo docx (no search, no clarifying question); a real owned file
    is created."""
    create_msg = _create_call(
        {
            "title": "Пример документа",
            "blocks": [{"type": "paragraph", "text": "Это демонстрационный документ."}],
        },
        output_format="docx",
    )
    final_msg = {"content": "Документ готов: Пример документа.docx"}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "сгенерируй любой документ"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert [c["name"] for c in data["tool_calls"]] == ["create_document"]
    assert data["tool_calls"][0]["arguments"]["output_format"] == "docx"

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    assert result["file_type"] == "docx"
    created = _get_document(result["document_id"])
    assert created is not None
    assert created.user_id == user_id
    assert Path(created.filepath).is_file()


def test_agent_creates_example_document_with_random_data(client, monkeypatch, user_id):
    """'создай пример с рандомными данными' -> create_document with clearly
    fictional demo data; no search and no clarifying questions."""
    create_msg = _create_call(
        {
            "title": "Пример трудового договора",
            "blocks": [
                {"type": "heading", "level": 1, "text": "Стороны"},
                {
                    "type": "paragraph",
                    "text": "Работодатель: ООО «Альфа». Работник: Иванов Иван Иванович, инженер.",
                },
                {"type": "paragraph", "text": "Документ носит демонстрационный характер."},
            ],
        }
    )
    final_msg = {"content": "Пример договора создан с демонстрационными данными."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": "создай пример документа с рандомными данными"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert [c["name"] for c in data["tool_calls"]] == ["create_document"]
    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    created = _get_document(result["document_id"])
    assert created is not None
    assert created.user_id == user_id
    assert "Иванов" in created.content


def test_agent_answers_salary_question_without_create(client, monkeypatch):
    """'сколько зарплата у Сергея' -> search -> read -> answer; NO
    create_document: answering a question is not creating a file."""
    marker = f"SAL{uuid.uuid4().hex[:6]}"
    text = (f"Кадровые данные {marker}. Зарплата Сергея Юрьевича — 180000 рублей.") * 20
    resp = _upload(client, "salary.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()[0]["id"]

    search_msg = _tool_call_message("search_documents", {"query": f"зарплата Сергея {marker}"})
    read_msg = _tool_call_message("read_document", {"document_id": document_id})
    final_msg = {"content": "Зарплата Сергея Юрьевича — 180000 рублей в месяц."}
    _scripted_functions(
        monkeypatch, [(search_msg, "s1"), (read_msg, "s2"), (final_msg, None)]
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": f"сколько зарплата у Сергея {marker}"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["answer"] == "Зарплата Сергея Юрьевича — 180000 рублей в месяц."
    assert [c["name"] for c in data["tool_calls"]] == ["search_documents", "read_document"]
    assert not any(c["name"] == "create_document" for c in data["tool_calls"])


def test_agent_answers_general_question_without_create(client, monkeypatch):
    """'расскажи кратко про налоговый кодекс' -> direct answer, no tools,
    and above all NO create_document: general knowledge is not a file request."""
    final_msg = {"content": "Налоговый кодекс РФ — основной акт налогового законодательства."}
    _scripted_functions(monkeypatch, [(final_msg, None)])

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": "расскажи кратко про налоговый кодекс"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    assert data["tool_calls"] == []
    assert data["tool_results"] == []
    assert "Налоговый кодекс" in data["answer"]


def test_agent_generates_pdf(client, monkeypatch, user_id):
    """'создай pdf': the model passes output_format='pdf' and a real PDF
    document is created and registered for the user."""
    create_msg = _create_call(
        {"title": "Договор", "blocks": [{"type": "paragraph", "text": "Условия PDF-договора."}]},
        output_format="pdf",
    )
    final_msg = {"content": "PDF готов."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай pdf документ"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    assert data["answer"] == final_msg["content"]

    from app.database.session import SessionLocal
    from app.models.document import Document

    db = SessionLocal()
    try:
        doc = (
            db.query(Document)
            .filter(Document.user_id == user_id, Document.file_type == "pdf")
            .first()
        )
    finally:
        db.close()
    assert doc is not None
    assert doc.original_filename.endswith(".pdf")
    from pathlib import Path

    assert Path(doc.filepath).read_bytes().startswith(b"%PDF")


def test_create_document_from_markdown_content(client, monkeypatch, user_id):
    """create_document accepts the document body as a Markdown `content` string
    and the backend parses it into the structured spec (headings/paragraphs)."""
    create_msg = _tool_call_message(
        "create_document",
        {
            "title": "Справка",
            "content": (
                "# Справка\n\n"
                "## Общие сведения\n\n"
                "ФИО: Иванов Иван Иванович\n"
                "Должность: инженер\n\n"
                "## Список\n\n"
                "- первый\n"
                "- второй\n"
            ),
            "output_format": "docx",
        },
    )
    final_msg = {"content": "готово"}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай"})
    assert resp.status_code == 200, resp.text
    result = json.loads(resp.json()["tool_results"][0]["content"])
    assert result["success"] is True

    doc = _get_document(result["document_id"])
    assert doc is not None
    assert "Иванов" in doc.content
    assert "первый" in doc.content


def test_markdown_to_spec_parses_structure():
    """Markdown maps to the right block types and heading levels."""
    from app.services.markdown_spec import markdown_to_spec

    spec = markdown_to_spec(
        "# Заголовок\n\n"
        "Текст абзаца\n\n"
        "## Подраздел\n\n"
        "- один\n"
        "- два\n\n"
        "| Имя | Возраст |\n"
        "|---|---|\n"
        "| Аня | 30 |\n",
        title="Явный заголовок",
    )
    assert spec.title == "Явный заголовок"
    types = [type(b).__name__ for b in spec.blocks]
    assert types[0].startswith("Heading")
    assert types[1].startswith("Paragraph")
    assert types[3].startswith("List")
    assert types[4].startswith("Table")
    assert spec.blocks[4].rows == [["Аня", "30"]]


def test_create_document_md_and_txt(client, monkeypatch):
    """create_document with output_format 'md'/'txt' saves a real downloadable file."""
    for fmt in ("md", "txt"):
        create_msg = _tool_call_message(
            "create_document",
            {
                "title": "Заметка",
                "content": "# Заметка\n\nПростой текст для проверки.",
                "output_format": fmt,
            },
        )
        final_msg = {"content": "готово"}
        _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

        resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай заметку"})
        assert resp.status_code == 200, resp.text
        result = json.loads(resp.json()["tool_results"][0]["content"])
        assert result["success"] is True, result
        assert result["file_type"] == fmt
        assert result["filename"].endswith(f".{fmt}")

        doc = _get_document(result["document_id"])
        assert doc is not None
        assert "Простой текст для проверки" in doc.content
        if fmt == "md":
            assert doc.content.startswith("# Заметка")


# --- agent memory + realtime streaming ---------------------------------------


def _collect_events(question: str, user_id: int, monkeypatch, script):
    """Drive run_agent_stream with scripted GigaChat turns and return events."""
    _scripted_functions(monkeypatch, script)
    from app.schemas.agent import AgentRequest

    events = list(agent_service.run_agent_stream(AgentRequest(question=question), user_id=user_id))
    return events


def test_agent_stream_emits_realtime_events(client, monkeypatch, user_id):
    """search/read/create each emit running+completed steps, then document_created
    and final — and NO chain-of-thought is ever emitted (M)."""
    search_msg = _tool_call_message("search_documents", {"query": "шаблон"})
    read_msg = _tool_call_message("read_document", {"document_id": 1})
    create_msg = _tool_call_message(
        "create_document",
        {"title": "Договор", "content": "# Договор\n\nТекст.", "output_format": "docx"},
    )
    final_msg = {"content": "Готово."}
    events = _collect_events(
        "создай договор",
        user_id,
        monkeypatch,
        [(search_msg, "s"), (read_msg, "s"), (create_msg, "s"), (final_msg, None)],
    )

    types = [e["type"] for e in events]
    assert "final" in types
    # Every emitted event is a safe action log event, never a thought/reasoning.
    for e in events:
        assert e["type"] in {"agent_step", "document_created", "final"}
        assert set(e.keys()) <= {
            "type", "step_id", "status", "tool", "message",
            "content", "sources", "chat_id", "document_id", "filename", "download_url",
        }

    steps = [e for e in events if e["type"] == "agent_step"]
    tools_seen = [s["tool"] for s in steps if s["status"] == "running"]
    assert "search_documents" in tools_seen
    assert "read_document" in tools_seen
    assert "create_document" in tools_seen

    created = [e for e in events if e["type"] == "document_created"]
    assert created, "document_created event expected for a successful create"
    assert isinstance(created[0]["document_id"], int)
    assert created[0]["document_id"] > 0
    assert created[0]["download_url"].endswith("/file")


def test_agent_step_error_event_on_tool_failure(client, monkeypatch, user_id):
    """A failing tool call surfaces as a separate error step event (L)."""
    bad_read = _tool_call_message("read_document", {"document_id": 999999})
    final_msg = {"content": "не смог прочитать"}
    events = _collect_events(
        "прочитай документ 999999",
        user_id,
        monkeypatch,
        [(bad_read, "s"), (final_msg, None)],
    )
    steps = [e for e in events if e["type"] == "agent_step"]
    assert any(s["tool"] == "read_document" and s["status"] == "error" for s in steps)


def test_agent_memory_restores_context_across_turns(client, monkeypatch, user_id):
    """The second turn receives the prior conversation + persisted task state (A, E)."""
    create_msg = _tool_call_message(
        "create_document",
        {"title": "Договор", "content": "# Договор\n\nТекст.", "output_format": "docx"},
    )
    final1 = {"content": "создал"}
    final2 = {"content": "ok"}

    def fake1(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        return (create_msg, "s") if messages[-1]["role"] == "user" and "создай" in messages[-1]["content"] else (final1, None)

    monkeypatch.setattr(gemini, "chat_with_functions", fake1)
    events1 = list(agent_service.run_agent_stream(_req("создай договор"), user_id=user_id))
    chat_id = next(e["chat_id"] for e in events1 if e["type"] == "final")

    calls2: list = []

    def fake2(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        calls2.append(list(messages))
        return final2, None

    monkeypatch.setattr(gemini, "chat_with_functions", fake2)
    list(agent_service.run_agent_stream(_req("что с документом", chat_id=chat_id), user_id=user_id))

    assert calls2, "second turn should have called the model"
    last_messages = calls2[0]
    roles = [m["role"] for m in last_messages]
    assert "user" in roles and roles.count("user") >= 2
    assert any(
        m.get("role") == "system" and "Контекст задачи" in m["content"]
        for m in last_messages
    )


def test_new_chat_has_clean_context(client, monkeypatch, user_id):
    """A brand-new chat starts with no injected task/document context (D)."""
    calls: list = []

    def fake(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        calls.append(list(messages))
        return {"content": "привет"}, None

    monkeypatch.setattr(gemini, "chat_with_functions", fake)
    list(agent_service.run_agent_stream(_req("привет"), user_id=user_id))
    assert calls
    roles = [m["role"] for m in calls[0]]
    assert roles == ["system", "user"]


def test_search_result_carries_normalized_metadata(client, monkeypatch, user_id):
    """search_documents results expose id/name/type/size/date metadata (C)."""
    from app.database.session import SessionLocal
    from app.models.document import Document

    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename="stored.docx",
            original_filename="Doc_алексей.docx",
            file_type="docx",
            file_size=1234,
            filepath="/tmp/x.docx",
            content="данные сотрудника",
        )
        db.add(doc)
        db.commit()
        doc_id = doc.id
    finally:
        db.close()

    class _Src:
        document_id = doc_id
        filename = "Doc_алексей.docx"
        score = 0.9
        text = "данные сотрудника"

    class _Chunk:
        source = _Src()

    monkeypatch.setattr(
        "app.services.agent.retrieve_context",
        lambda *a, **k: [_Chunk()],
    )

    hits = agent_service._search_documents(user_id, "алексей", None)
    assert hits
    hit = hits[0]
    assert hit["document_id"] == doc_id
    assert hit["filename"] == "Doc_алексей.docx"
    assert hit["type"] == "docx"
    assert hit["file_size"] == 1234
    assert hit["owner_id"] == user_id
    assert "created_at" in hit
    assert hit["content_available"] is True


def _req(question: str, chat_id: int | None = None):
    from app.schemas.agent import AgentRequest

    return AgentRequest(question=question, chat_id=chat_id)


# --------------------------------------------------------------------------- #
# Explicit (pinned) context: edit_document must use the attached document id
# --------------------------------------------------------------------------- #


def _make_pdf_with_text(text: str) -> bytes:
    import fitz
    import io

    doc = fitz.open()
    doc.new_page()
    doc[0].insert_text(fitz.Point(50, 50), text)
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    content = buf.getvalue()
    doc.close()
    return content


def _seed_pdf(
    user_id: int, data: bytes, original_filename: str = "manual.pdf"
) -> tuple[int, Path]:
    from app.database.session import SessionLocal
    from app.models.document import Document

    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}.pdf"
    path = upload_dir / stored
    path.write_bytes(data)

    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename=stored,
            original_filename=original_filename,
            file_type="pdf",
            file_size=len(data),
            filepath=str(path),
            content="pdf seed",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc.id, path
    finally:
        db.close()


def test_agent_uses_pinned_pdf_context_for_edit(client, monkeypatch, user_id):
    """A single pinned document must become the edit_document target without the
    model (or the user) ever supplying its id. Even when the scripted model
    omits file_id, the agent resolves it from context_document_ids."""
    pdf_id, pdf_path = _seed_pdf(user_id, _make_pdf_with_text("LXSHOW user manual."))
    original = pdf_path.read_bytes()

    # The model calls edit_document WITHOUT file_id — the agent must inject it.
    edit_msg = _tool_call_message("edit_document", {"instruction": "Переведи на русский, уберите LXSHOW"})
    final_msg = {"content": "Переведено."}
    _scripted_functions(monkeypatch, [(edit_msg, "s"), (final_msg, None)])
    # Keep the edit deterministic (no real LLM for the block rewrite).
    monkeypatch.setattr(
        "app.services.document_edit._request_edits",
        lambda blocks, instruction: [f"[EDITED] {b}" for b in blocks],
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={
            "question": (
                "Переведи руководство пользователя на русский язык, сохраняя все "
                "картинки и верстку. Убери все упоминания LXSHOW. Отдай переведённый "
                "документ в формате pdf ОБЯЗАТЕЛЬНО"
            ),
            "context_document_ids": [pdf_id],
        },
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    # Exactly one tool ran: edit_document on the pinned PDF. No search_documents
    # (explicit context > RAG), and the model was never asked for an id.
    tool_names = {c["name"] for c in data["tool_calls"]}
    assert tool_names == {"edit_document"}
    assert data["tool_calls"][0]["arguments"].get("file_id") == pdf_id

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    assert result["file_type"] == "pdf"
    assert result["source_file_id"] == pdf_id
    # Original on disk is untouched.
    assert pdf_path.read_bytes() == original

    # CONFIRMATION QUALITY: the edit result cites the NEW file and states the
    # original is unchanged, with a real download url.
    assert result["download_url"].startswith(f"{API_PREFIX}/documents/")
    assert result["confirmation"]
    assert result["filename"] in result["confirmation"]
    assert "не изменён" in result["confirmation"].lower()


def test_agent_ignores_unpinned_docs_when_one_pinned(client, monkeypatch, user_id):
    """A pinned PDF must be edited directly; an unrelated library doc (Savvaland)
    must never be searched or read, and must not appear anywhere in the trace."""
    pdf_id, pdf_path = _seed_pdf(user_id, _make_pdf_with_text("LXSHOW manual."))
    decoy_id = _insert_document(user_id, "Savvaland.txt", "unrelated text", file_type="txt")
    original_pdf = pdf_path.read_bytes()

    # Stub the block rewrite so no real LLM is needed.
    monkeypatch.setattr(
        "app.services.document_edit._request_edits",
        lambda blocks, instruction: [f"[EDITED] {b}" for b in blocks],
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={
            "question": (
                "Переведи этот документ на русский язык. Убери LXSHOW. Верни PDF."
            ),
            "context_document_ids": [pdf_id],
        },
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    tool_names = {c["name"] for c in data["tool_calls"]}
    assert tool_names == {"edit_document"}, tool_names
    assert data["tool_calls"][0]["arguments"].get("file_id") == pdf_id
    # The decoy is never referenced.
    assert str(decoy_id) not in str(data)

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    assert result["source_file_id"] == pdf_id
    assert pdf_path.read_bytes() == original_pdf


def test_agent_resolves_named_pinned_doc_among_many(client, monkeypatch, user_id):
    """With several pinned docs, a request that names one resolves to it."""
    manual_id, _ = _seed_pdf(user_id, _make_pdf_with_text("manual body"))
    source_id, _ = _seed_pdf(
        user_id, _make_pdf_with_text("source body"), original_filename="source.pdf"
    )

    edit_msg = _tool_call_message("edit_document", {"instruction": "Переведи manual.pdf"})
    final_msg = {"content": "Переведено."}
    _scripted_functions(monkeypatch, [(edit_msg, "s"), (final_msg, None)])
    monkeypatch.setattr(
        "app.services.document_edit._request_edits",
        lambda blocks, instruction: [f"[EDITED] {b}" for b in blocks],
    )

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={
            "question": "Переведи manual.pdf на русский",
            "context_document_ids": [manual_id, source_id],
        },
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["tool_calls"][0]["arguments"].get("file_id") == manual_id


# ------------------------------------------------- create_document placeholder gate


def test_create_document_rejects_critical_placeholders(client, monkeypatch):
    """A "finished" document that still contains unfilled fields ({{...}},
    [дата], [подписи], TODO) must NOT be silently created: create_document
    returns success:false with the offending placeholders listed, so the
    model asks for the missing data instead of claiming the file is ready."""
    create_msg = _create_call(
        {
            "title": "Договор",
            "blocks": [
                {"type": "paragraph", "text": "Дата подписания: [дата]."},
                {"type": "paragraph", "text": "Сумма: {{SUM}}."},
            ],
        }
    )
    final_msg = {"content": "Не хватает данных."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай договор"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is False
    assert result["error_type"] == "DocumentIncompleteError"
    assert "[дата]" in result["error"]
    assert "{{SUM}}" in result["error"]
    assert set(result["placeholders"]) == {"[дата]", "{{SUM}}"}


def test_create_document_allows_placeholders_for_template_request(client, monkeypatch, user_id):
    """'по шаблону' requests legitimately keep placeholder slots, so the same
    spec IS created when the user explicitly asked for a template."""
    create_msg = _create_call(
        {
            "title": "Типовой договор",
            "blocks": [
                {"type": "paragraph", "text": "Дата: [дата]."},
            ],
        }
    )
    final_msg = {"content": "Шаблон готов."}
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(
        f"{API_PREFIX}/agent",
        json={"question": "сделай шаблон договора"},
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    created = _get_document(result["document_id"])
    assert created is not None
    assert created.user_id == user_id
    assert "[дата]" in created.content


def test_search_self_corrects_zero_hit_query(client, monkeypatch, user_id):
    """When the model's search query returns nothing, the agent rewrites it
    (strip pleading verbs, drop function words) and searches again instead of
    answering 'not found'. The recovered hit is labelled reformulated_query."""
    from app.models.document import Document
    from app.database.session import SessionLocal

    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename="salary.txt",
            original_filename="salary.txt",
            file_type="txt",
            file_size=64,
            filepath="/data/uploads/salary.txt",
            content="зарплата Сергея Юрьевича — 180000 рублей",
        )
        db.add(doc)
        db.commit()
        doc_id = doc.id
    finally:
        db.close()

    class _Src:
        document_id = doc_id
        filename = "salary.txt"
        score = 0.9
        text = "зарплата Сергея Юрьевича — 180000 рублей"

    class _Chunk:
        source = _Src()

    original = "найди мне данные про зарплату Сергея"
    rewritten = "данные про зарплату сергея"  # stripped of 'найди мне'

    def fake_retrieve(**kwargs):
        question = (kwargs.get("question") or "").lower()
        if question == original.lower():
            return []
        if question == rewritten:
            return [_Chunk()]
        return []

    monkeypatch.setattr("app.services.agent.retrieve_context", fake_retrieve)

    tool_msg = _tool_call_message("search_documents", {"query": original})
    final_msg = {"content": "Зарплата Сергея Юрьевича — 180000 рублей."}
    _scripted_functions(monkeypatch, [(tool_msg, "s"), (final_msg, None)])

    resp = client.post(
        f"{API_PREFIX}/agent", json={"question": f"какая зарплата у {doc_id}"}
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()

    result = json.loads(data["tool_results"][0]["content"])
    assert result, "the reformulated search must recover a hit"
    hit = result[0]
    assert hit["document_id"] == doc_id
    assert hit["filename"] == "salary.txt"
    assert hit["reformulated_query"] == rewritten
    assert "180000" in hit["snippet"]


def test_empty_final_answer_falls_back_to_confirmation(client, monkeypatch, user_id):
    """CONFIRMATION QUALITY: if the model creates a document but returns an
    empty final answer, the backend replies with the factual confirmation from
    the tool result instead of an empty box."""
    create_msg = _create_call(
        {
            "title": "Трудовой договор",
            "blocks": [{"type": "paragraph", "text": "Стороны заключили договор."}],
        }
    )
    final_msg = {"content": ""}  # model silently returns nothing
    _scripted_functions(monkeypatch, [(create_msg, "s"), (final_msg, None)])

    resp = client.post(f"{API_PREFIX}/agent", json={"question": "создай договор"})
    assert resp.status_code == 200, resp.text
    data = resp.json()

    result = json.loads(data["tool_results"][0]["content"])
    assert result["success"] is True
    # The answer is the factual confirmation (real filename + format).
    assert data["answer"]
    assert result["filename"] in data["answer"]
    assert "DOCX" in data["answer"]


