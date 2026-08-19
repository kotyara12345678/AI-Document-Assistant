"""Agent policy tests: the agent should decide on its own when to retrieve,
rephrase or read more, admit lack of data, extract before generating, and wrap
copyable documents in a single fenced code block.

The LLM is mocked at the ``chat_with_functions`` boundary (same pattern as
test_agent.py / test_agent_discovery.py). Behavioural tests drive the REAL
retrieval pipeline (PostgreSQL FTS + Qdrant) and the ``/api/agent`` endpoint.
"""

import io
import json
import os
import uuid
import zipfile

import docx
from fastapi.testclient import TestClient

from app.database.session import SessionLocal
from app.models.document import Document
from app.services import gemini
from app.services.agent import SYSTEM_INSTRUCTION

API_PREFIX = "/api"


def _make_odt_bytes(text: str) -> bytes:
    paragraphs = "\n".join(f"<text:p>{line}</text:p>" for line in text.splitlines())
    content_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
  office:version="1.2">
  <office:body><office:text>{paragraphs}</office:text></office:body>
</office:document-content>"""
    manifest = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<manifest:manifest '
        'xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" '
        'manifest:version="1.2">'
        '<manifest:file-entry manifest:full-path="/" manifest:version="1.2" '
        'manifest:media-type="application/vnd.oasis.opendocument.text"/>'
        '<manifest:file-entry manifest:full-path="content.xml" media-type="text/xml"/>'
        '<manifest:file-entry manifest:full-path="mimetype" '
        'media-type="application/vnd.oasis.opendocument.text"/>'
        "</manifest:manifest>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("mimetype", b"application/vnd.oasis.opendocument.text")
        zf.writestr("META-INF/manifest.xml", manifest.encode("utf-8"))
        zf.writestr("content.xml", content_xml.encode("utf-8"))
    return buf.getvalue()


def _upload(client: TestClient, filename: str, content: bytes) -> int:
    resp = client.post(
        f"{API_PREFIX}/documents/upload",
        files={"file": (filename, content)},
    )
    assert resp.status_code == 201, resp.text
    return resp.json()[0]["id"]


def _tool_call_message(name="search_documents", arguments=None):
    return {
        "role": "assistant",
        "content": None,
        "function_call": {"name": name, "arguments": arguments or {}},
    }


def _scripted_functions(monkeypatch, script):
    """Queue (message, state_id) pairs; record every outgoing request."""
    calls = []

    def fake(messages, functions=None, function_call="auto", functions_state_id=None, client=None, usage_hook=None):
        calls.append(
            {
                "messages": messages,
                "functions": functions,
                "function_call": function_call,
                "functions_state_id": functions_state_id,
            }
        )
        return script.pop(0)

    monkeypatch.setattr(gemini, "chat_with_functions", fake)
    return calls


# ---------------------------------------------------------------------------
# Prompt-guard tests (no DB needed) — pin the policy added for goals 5-9.
# ---------------------------------------------------------------------------


class TestSystemInstructionPolicy:
    def test_proactive_retrieval_without_explicit_hint(self):
        # Goal 1 (strengthened): search even without "find"/"document" wording.
        text = SYSTEM_INSTRUCTION.lower()
        assert "presumed to need retrieval" in text
        assert "search even if the user never says" in text

    def test_uses_metadata_only_when_needed(self):
        # Goal 5: metadata from tools, never invent, only when it helps.
        text = SYSTEM_INSTRUCTION.lower()
        assert "use the metadata the tools return" in text
        assert "never invent metadata" in text
        assert "mention it only when it directly helps" in text

    def test_document_body_goes_into_blocks_not_chat(self):
        # The full document must be delivered via the create_document 'content'
        # argument (Markdown), not as a code block in the chat reply.
        text = SYSTEM_INSTRUCTION.lower()
        assert "into the 'content'" in text
        assert "never wrap the generated document in a code block" in text
        assert "without 'content'" in text

    def test_avoids_code_block_for_plain_answers(self):
        # Goal 9: ordinary answers stay plain text.
        text = SYSTEM_INSTRUCTION.lower()
        assert "never wrap them in a code block" in text
        assert "use plain text" in text
        assert "ordinary (non-document) answers" in text

    def test_document_generation_rule_mandates_create_document(self):
        # Document generation rule: create_document is REQUIRED and the model
        # must never claim generation is unavailable or substitute chat text.
        text = SYSTEM_INSTRUCTION.lower()
        assert "call create_document" in text
        assert "tool call is required" in text
        assert "never claim that document generation is unavailable" in text
        assert "never substitute writing the document in the chat" in text
        assert "put the entire document into" in text

    def test_confirmation_must_cite_real_tool_result(self):
        # CONFIRMATION QUALITY: after create/edit the reply must be based on the
        # actual tool result, never on invented file names or formats.
        text = SYSTEM_INSTRUCTION.lower()
        assert "confirmation quality" in text
        assert "based on the tool result" in text
        assert "real file name" in text
        assert "never invent a file name" in text
        assert "cannot provide" in text

    def test_edit_confirmation_states_original_unchanged(self):
        # After edit_document the model must say a NEW file was created and the
        # original document is unchanged.
        text = SYSTEM_INSTRUCTION.lower()
        assert "a new file was created" in text
        assert "original document is" in text
        assert "не измен" in text

    def test_anti_fabrication_rule_pinned(self):
        # Critical regression: the model must never claim tools were called,
        # results returned, documents read, files created or links provided
        # unless the real tool result in the conversation proves it.
        text = SYSTEM_INSTRUCTION.lower()
        assert "anti-fabrication" in text
        assert "never claim that you called a tool" in text
        assert "unless the actual tool result" in text
        assert "tool was not called, failed" in text
        assert "never invent tool calls" in text
        assert "download_url" in text
        assert "success: true tool result" in text


# ---------------------------------------------------------------------------
# Behavioural tests (real retrieval + mocked LLM).
# ---------------------------------------------------------------------------


class TestAgentBehaviour:
    def test_proactive_search_without_keyword(self, client, monkeypatch):
        """Goal 1: a factual question with no 'find'/'document' wording still
        retrieves and reads before answering."""
        marker = uuid.uuid4().hex[:8]
        text = (
            f"Сотрудник Иван Петров {marker} получает оклад 180000 рублей." * 5
        )
        _upload(client, f"emp_{marker}.txt", text.encode("utf-8"))

        script = [
            (_tool_call_message("search_documents", {"query": f"Иван Петров {marker} оклад"}), "s1"),
            (_tool_call_message("read_document", {"document_id": 1, "offset": 0}), "s2"),
            ({"content": f"Оклад Ивана Петрова {marker} — 180000 рублей."}, None),
        ]
        _scripted_functions(monkeypatch, script)

        resp = client.post(f"{API_PREFIX}/agent", json={"question": f"Сколько получает Иван Петров {marker}?"})
        assert resp.status_code == 200, resp.text
        data = resp.json()
        names = [c["name"] for c in data["tool_calls"]]
        assert names == ["search_documents", "read_document"]

    def test_runs_multiple_retrievals_for_complex_question(self, client, monkeypatch):
        """Goal 2: a complex question is resolved through several searches /
        rephrasings and reads, not a single shot."""
        marker = uuid.uuid4().hex[:8]
        text = (
            f"Проект {marker}: бюджет 1 200 000 рублей, срок сдачи 30 дней." * 5
        )
        _upload(client, f"proj_{marker}.txt", text.encode("utf-8"))

        script = [
            (_tool_call_message("search_documents", {"query": f"бюджет проекта {marker}"}), "s1"),
            (_tool_call_message("search_documents", {"query": f"{marker} срок сдачи финансы смета"}), "s2"),
            (_tool_call_message("read_document", {"document_id": 1, "offset": 0}), "s3"),
            ({"content": f"Бюджет проекта {marker} — 1 200 000 рублей, срок — 30 дней."}, None),
        ]
        _scripted_functions(monkeypatch, script)

        resp = client.post(
            f"{API_PREFIX}/agent",
            json={"question": f"Какой бюджет и срок у проекта {marker}?"},
        )
        assert resp.status_code == 200, resp.text
        data = resp.json()
        names = [c["name"] for c in data["tool_calls"]]
        assert names.count("search_documents") == 2
        assert names.count("read_document") == 1

    def test_reports_insufficient_information(self, client, monkeypatch):
        """Goal 3: when retrieval yields nothing relevant, answer honestly
        instead of inventing."""
        script = [
            (_tool_call_message("search_documents", {"query": "телефон заказчика контракта"}), "s1"),
            ({"content": "В ваших документах нет информации о телефоне заказчика."}, None),
        ]
        _scripted_functions(monkeypatch, script)

        resp = client.post(
            f"{API_PREFIX}/agent",
            json={"question": "Какой телефон у заказчика в договоре?"},
        )
        assert resp.status_code == 200, resp.text
        data = resp.json()
        names = [c["name"] for c in data["tool_calls"]]
        assert names == ["search_documents"]
        assert "нет" in data["answer"].lower()

    def test_plain_question_needs_no_retrieval(self, client, monkeypatch):
        """Goal 9: a general, document-independent question gets a direct answer
        with no tool calls."""
        _scripted_functions(monkeypatch, [({"content": "Здравствуйте! Чем могу помочь?"}, None)])
        resp = client.post(f"{API_PREFIX}/agent", json={"question": "Привет, как дела?"})
        assert resp.status_code == 200, resp.text
        data = resp.json()
        assert data["tool_calls"] == []
        assert data["answer"]

    def test_real_scenario_doc_aleksey_contract(self, client, monkeypatch):
        """REGRESSION: 'Используй Doc_алексей и составь трудовой договор.
        Сгенерируй готовый документ по шаблону' must retrieve, read and call
        create_document — never refuse that generation is unavailable."""
        source = (
            "ТРУДОВОЙ ДОГОВОР. Работник Иван Иванов. Оклад 200000 рублей. "
            "Должность инженер." * 3
        )
        doc_id = _upload(client, "Doc_алексей.txt", source.encode("utf-8"))

        script = [
            (_tool_call_message("search_documents", {"query": "Doc_алексей трудовой договор"}), "s1"),
            (_tool_call_message("read_document", {"document_id": doc_id, "offset": 0}), "s2"),
            (
                _tool_call_message(
                    "create_document",
                    {
                        "document_spec": {
                            "title": "Трудовой договор",
                            "blocks": [
                                {"type": "heading", "level": 1, "text": "Трудовой договор"},
                                {"type": "paragraph", "text": "Работник: Иван Иванов."},
                                {"type": "paragraph", "text": "Оклад: 200000 рублей."},
                            ],
                        },
                        "output_format": "docx",
                    },
                ),
                "s3",
            ),
            ({"content": "Договор готов и сохранён."}, None),
        ]
        _scripted_functions(monkeypatch, script)

        resp = client.post(
            f"{API_PREFIX}/agent",
            json={
                "question": "Используй Doc_алексей и составь трудовой договор. "
                "Сгенерируй готовый документ по шаблону."
            },
        )
        assert resp.status_code == 200, resp.text
        data = resp.json()
        names = [c["name"] for c in data["tool_calls"]]
        assert names == ["search_documents", "read_document", "create_document"], names
        created = [r for r in data["tool_results"] if r["name"] == "create_document"]
        assert created, "create_document was not called"
        payload = json.loads(created[-1]["content"])
        assert payload.get("success") is True, "create_document did not succeed"
        doc_id = payload["document_id"]

        # The actual generated file must exist in the DB and on disk — we do not
        # trust the model's "document is ready" text, only the real tool result.
        db = SessionLocal()
        try:
            doc = db.query(Document).filter(Document.id == doc_id).first()
        finally:
            db.close()
        assert doc is not None, f"generated document {doc_id} was not persisted"
        assert os.path.exists(doc.filepath), f"generated file missing on disk: {doc.filepath}"

        # The generated document must be surfaced to the UI so the user can fetch it.
        assert any(
            c["document_id"] == doc_id for c in data.get("created_documents", [])
        ), "created document was not returned to the frontend"

    def test_create_document_renders_full_content_into_docx(self, client, monkeypatch):
        """REGRESSION: the generated .docx must contain the WHOLE document —
        headings, paragraphs, lists, tables and Cyrillic text — not just the
        title. The model must put the body into document_spec.blocks."""
        script = [
            (
                _tool_call_message(
                    "create_document",
                    {
                        "document_spec": {
                            "title": "Трудовой договор",
                            "blocks": [
                                {"type": "heading", "level": 1, "text": "Предмет договора"},
                                {
                                    "type": "paragraph",
                                    "text": "Работодатель обязуется предоставить работу по должности.",
                                },
                                {"type": "heading", "level": 1, "text": "Условия оплаты"},
                                {
                                    "type": "paragraph",
                                    "text": "Оклад составляет 200000 рублей ежемесячно.",
                                },
                                {
                                    "type": "list",
                                    "ordered": False,
                                    "items": [
                                        "Медицинское страхование",
                                        "Отпуск 28 календарных дней",
                                    ],
                                },
                                {
                                    "type": "table",
                                    "headers": ["Позиция", "Значение"],
                                    "rows": [["Должность", "Инженер"], ["Город", "Москва"]],
                                },
                            ],
                        },
                        "output_format": "docx",
                    },
                ),
                "s1",
            ),
            ({"content": "Документ «Трудовой договор» создан."}, None),
        ]
        _scripted_functions(monkeypatch, script)

        resp = client.post(f"{API_PREFIX}/agent", json={"question": "Составь трудовой договор."})
        assert resp.status_code == 200, resp.text
        data = resp.json()
        assert [c["name"] for c in data["tool_calls"]] == ["create_document"]
        created = [r for r in data["tool_results"] if r["name"] == "create_document"]
        assert created, "create_document was not called"
        payload = json.loads(created[-1]["content"])
        assert payload.get("success") is True, "create_document did not succeed"
        doc_id = payload["document_id"]

        db = SessionLocal()
        try:
            doc = db.query(Document).filter(Document.id == doc_id).first()
        finally:
            db.close()
        assert doc is not None, f"generated document {doc_id} was not persisted"
        assert os.path.exists(doc.filepath), f"generated file missing on disk: {doc.filepath}"

        # Open the real .docx and confirm the entire body is present.
        document = docx.Document(doc.filepath)
        full_text = "\n".join(p.text for p in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                full_text += "\n" + " | ".join(cell.text for cell in row.cells)

        for expected in [
            "Трудовой договор",
            "Предмет договора",
            "Работодатель обязуется предоставить работу по должности.",
            "Условия оплаты",
            "Оклад составляет 200000 рублей ежемесячно.",
            "Медицинское страхование",
            "Отпуск 28 календарных дней",
            "Должность",
            "Инженер",
            "Москва",
        ]:
            assert expected in full_text, f"missing from generated docx: {expected}"
