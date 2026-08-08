"""E2E tests for the document → chunks → embeddings → Qdrant → search pipeline.

Run inside the backend container against the live stack:
    docker compose exec backend pytest -q
"""

import io
import uuid

import httpx
import pytest
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.core.config import settings
from app.database.session import SessionLocal
from app.models.document import Document
from app.vector.client import delete_document_vectors, get_qdrant_client

BASE_URL = "http://localhost:8000"
API_PREFIX = settings.API_PREFIX
USER_ID = 1  # matches get_current_user_id() stub


@pytest.fixture()
def client():
    with httpx.Client(base_url=BASE_URL, timeout=120) as c:
        yield c


@pytest.fixture()
def db_session():
    db: Session = SessionLocal()
    yield db
    db.close()


@pytest.fixture(autouse=True)
def _clean_qdrant():
    """Isolate each test from points left by previous runs."""
    client = get_qdrant_client()
    try:
        client.delete_collection(settings.QDRANT_COLLECTION)
    except Exception:
        pass
    yield


def _make_pdf_bytes() -> bytes:
    """Build a minimal valid PDF containing searchable text."""
    text = "This is a test PDF document with searchable budget information 1500000."
    content = b"BT /F1 24 Tf 100 700 Td (" + text.encode("latin-1") + b") Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = bytearray()
    out += b"%PDF-1.4\n"
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj\n" + obj + b"\nendobj\n"

    xref_pos = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
    out += b"startxref\n" + str(xref_pos).encode() + b"\n%%EOF\n"
    return bytes(out)


def _make_docx_bytes() -> bytes:
    doc = DocxDocument()
    doc.add_paragraph("Тестовый DOCX документ для проверки индексации.")
    doc.add_paragraph("Ожидаемый бюджет проекта составляет 500000 рублей.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload(client: httpx.Client, filename: str, content: bytes):
    files = {"file": (filename, content)}
    return client.post(f"{API_PREFIX}/documents/upload", files=files)


def _document_ids_in_db(db: Session) -> set[int]:
    return {row[0] for row in db.query(Document.id).filter(Document.user_id == USER_ID).all()}


def _qdrant_point_ids(document_id: int) -> list:
    client = get_qdrant_client()
    try:
        res = client.scroll(
            collection_name=settings.QDRANT_COLLECTION,
            limit=1000,
            scroll_filter={"must": [{"key": "document_id", "match": {"value": document_id}}]},
            with_payload=False,
        )
        return [p.id for p in res[0]]
    except Exception:
        return []


def test_upload_txt(client, db_session):
    content = "Просто тестовый текст в формате TXT с полезной информацией для поиска."
    resp = _upload(client, "sample.txt", content.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    data = resp.json()
    assert data["file_type"] == "txt"
    assert data["content_length"] == len(content)
    assert data["id"] in _document_ids_in_db(db_session)


def test_upload_pdf(client, db_session):
    content = _make_pdf_bytes()
    resp = _upload(client, "sample.pdf", content)
    assert resp.status_code == 201, resp.text
    data = resp.json()
    assert data["file_type"] == "pdf"
    assert data["file_size"] == len(content)
    assert data["id"] in _document_ids_in_db(db_session)


def test_upload_docx(client, db_session):
    content = _make_docx_bytes()
    resp = _upload(client, "sample.docx", content)
    assert resp.status_code == 201, resp.text
    data = resp.json()
    assert data["file_type"] == "docx"
    assert data["id"] in _document_ids_in_db(db_session)


def test_upload_invalid_extension(client):
    resp = _upload(client, "virus.exe", b"MZ....")
    assert resp.status_code == 400


def test_upload_empty_file(client):
    resp = _upload(client, "empty.txt", b"")
    assert resp.status_code == 400


def test_upload_triggers_indexing(client):
    marker = f"OBS{uuid.uuid4().hex[:6]}"
    text = (
        f"Финансовое планирование астрономической обсерватории {marker}. "
        "Годовой бюджет телескопа составляет 2500000 рублей. "
        "Основные статьи расходов: оптика, автоматизация, персонал. "
    ) * 20
    resp = _upload(client, "observatory.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()["id"]

    points = _qdrant_point_ids(document_id)
    assert len(points) > 0, "No Qdrant points found for the uploaded document"


def test_search_returns_related_chunk(client):
    marker = f"KQ{uuid.uuid4().hex[:6]}"
    text = (
        f"Бюджет маркетингового отдела компании по производству квадрокоптеров {marker}. "
        "Рекламный бюджет составляет 900000 рублей на квартал. "
        "Из них на контекстную рекламу выделено 300000 рублей. "
    ) * 20
    resp = _upload(client, "quadcopter_budget.txt", text.encode("utf-8"))
    assert resp.status_code == 201, resp.text
    document_id = resp.json()["id"]

    search_resp = client.post(
        f"{API_PREFIX}/search",
        json={"query": f"рекламный бюджет квадрокоптеры {marker}", "limit": 5},
    )
    assert search_resp.status_code == 200, search_resp.text
    results = search_resp.json()["results"]
    assert results, "Search returned no results"

    top = results[0]
    assert top["score"] > 0.3, f"Score unexpectedly low: {top['score']}"
    assert top["document_id"] == document_id
    assert top["chunk_index"] >= 0
    assert top["filename"] == "quadcopter_budget.txt"
    assert "900000" in top["text"]


def test_manual_reindex_endpoint(client):
    resp = _upload(client, "reindex.txt", ("восстановление потерянной индексации роботов " * 50).encode("utf-8"))
    assert resp.status_code == 201
    document_id = resp.json()["id"]

    delete_document_vectors(document_id)
    assert _qdrant_point_ids(document_id) == []

    index_resp = client.post(f"{API_PREFIX}/documents/{document_id}/index")
    assert index_resp.status_code == 200, index_resp.text
    data = index_resp.json()
    assert data["chunks_indexed"] > 0
    assert len(_qdrant_point_ids(document_id)) == data["chunks_indexed"]


def test_reindex_missing_document_404(client):
    resp = client.post(f"{API_PREFIX}/documents/{uuid.uuid4().int % 100000000}/index")
    assert resp.status_code == 404


def test_delete_vectors(client):
    resp = _upload(client, "to_delete.txt", ("удаляемый документ " * 10).encode("utf-8"))
    assert resp.status_code == 201
    document_id = resp.json()["id"]
    assert len(_qdrant_point_ids(document_id)) > 0

    delete_document_vectors(document_id)
    assert _qdrant_point_ids(document_id) == []
