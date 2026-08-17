"""Regression tests for the LLM block-editing contract in ``document_edit``.

The GigaChat client is mocked so we exercise the parsing/retry logic of
``_request_edits`` directly, plus a full PDF translation integration path that
proves the original file is untouched and images/pages survive.
"""
import io
import json
import uuid
from pathlib import Path

import fitz
import pytest

from app.core.config import settings
from app.database.session import SessionLocal
from app.models.document import Document
from app.services import gemini
from app.services.document_edit import (
    _parse_edits_array,
    _request_edits,
    edit_document,
)
from app.services.errors import DocumentEditError


def _make_pdf_with_image(text1: str, text2: str) -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(fitz.Point(50, 50), text1)
    page.insert_text(fitz.Point(50, 80), text2)
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 24, 24), False)
    page.insert_image(fitz.Rect(300, 700, 360, 760), stream=pix.tobytes("png"))
    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    data = buf.getvalue()
    doc.close()
    return data


def _seed_pdf(user_id: int, data: bytes) -> tuple[int, Path]:
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}.pdf"
    path = upload_dir / stored
    path.write_bytes(data)
    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename=stored,
            original_filename="manual.pdf",
            file_type="pdf",
            file_size=len(data),
            filepath=str(path),
            content="pdf seed",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc.id, path
    finally:
        db.close()


def test_request_edits_accepts_plain_json_array(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: '["one","two"]')
    assert _request_edits(["a", "b"], "translate") == ["one", "two"]


def test_request_edits_accepts_fenced_json(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: '```json\n["one","two"]\n```')
    assert _request_edits(["a", "b"], "translate") == ["one", "two"]


def test_request_edits_handles_trailing_commas(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: '["one","two",]')
    assert _request_edits(["a", "b"], "translate") == ["one", "two"]


def test_request_edits_accepts_array_of_objects(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: '[{"text":"one"},{"text":"two"}]')
    assert _request_edits(["a", "b"], "translate") == ["one", "two"]


def test_request_edits_retries_on_invalid_then_succeeds(monkeypatch):
    calls = {"n": 0}

    def fake(prompt, system_instruction=None, **kwargs):
        calls["n"] += 1
        return "Sure! Here is the result: not-json" if calls["n"] == 1 else '["one","two"]'

    monkeypatch.setattr(gemini, "generate_answer", fake)
    assert _request_edits(["a", "b"], "translate") == ["one", "two"]
    assert calls["n"] == 2


def test_request_edits_rejects_invalid_after_retry(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: "no json here")
    with pytest.raises(DocumentEditError):
        _request_edits(["a", "b"], "translate")


def test_request_edits_rejects_length_mismatch(monkeypatch):
    monkeypatch.setattr(gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: '["only","one","extra"]')
    with pytest.raises(DocumentEditError):
        _request_edits(["a", "b"], "translate")


def test_request_edits_chunks_large_block_count(monkeypatch):
    """A large document must be split into chunks; the concatenated edits must
    exactly match the original block count and order."""

    def fake(prompt, system_instruction=None, **kwargs):
        # Echo every input block back unchanged so we can verify a strict 1:1
        # mapping (count + order) after the chunks are concatenated.
        blob = (system_instruction or "") + "\n" + (prompt or "")
        texts = []
        for line in blob.splitlines():
            idx, sep, text = line.partition("\t")
            if sep and idx.isdigit():
                texts.append(text)
        return json.dumps([{"text": t} for t in texts])

    monkeypatch.setattr(gemini, "generate_answer", fake)

    total = 75  # > EDIT_CHUNK_SIZE (30) -> 3 chunks
    blocks = [f"block {i} text" for i in range(total)]
    result = _request_edits(blocks, "translate to russian, remove LXSHOW")

    assert len(result) == total
    assert result == blocks


def test_parse_edits_array_rejects_truncated_json():
    # A mid-array truncation (e.g. from a token cap) must NOT be accepted.
    truncated = '[{"text": "one"}, {"text": "two"'
    assert _parse_edits_array(truncated, expected=2) is None


def test_request_edits_single_large_block(monkeypatch):
    # A single oversized block cannot be split; it must still be edited as one.
    big = "LXSHOW " * 1500
    monkeypatch.setattr(
        gemini, "generate_answer", lambda prompt, system_instruction=None, **kwargs: f'[{{"text": "{big.replace("LXSHOW", "X")}"}}]'
    )
    result = _request_edits([big], "remove LXSHOW")
    assert len(result) == 1
    assert "LXSHOW" not in result[0]


def test_request_edits_preserves_order(monkeypatch):
    blocks = [f"block-{i}" for i in range(12)]

    def fake(prompt, system_instruction=None, **kwargs):
        blob = (system_instruction or "") + "\n" + (prompt or "")
        texts = []
        for line in blob.splitlines():
            idx, sep, text = line.partition("\t")
            if sep and idx.isdigit():
                texts.append(text)
        return json.dumps([{"text": t} for t in texts])

    monkeypatch.setattr(gemini, "generate_answer", fake)
    assert _request_edits(blocks, "translate") == blocks


def test_pdf_large_translate_preserves_images_pages(monkeypatch, user_id):
    # Multi-page PDF with an image per page and many text blocks -> chunking is
    # exercised; images, page count and the original file must survive.
    doc = fitz.open()
    n_pages = 3
    for p in range(n_pages):
        page = doc.new_page()
        pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 24, 24), False)
        page.insert_image(fitz.Rect(300, 700, 360, 760), stream=pix.tobytes("png"))
        for i in range(15):
            page.insert_text(fitz.Point(50, 60 + i * 18), f"Para {p}-{i}. LXSHOW device info.")

    buf = io.BytesIO()
    doc.save(buf, garbage=4, deflate=True)
    content = buf.getvalue()
    doc.close()

    source_id, source_path = _seed_pdf(user_id, content)

    # Stub the structured edit step (the protocol the PDF path actually uses):
    # echo every item, removing LXSHOW. Exercises extraction -> chunk -> apply
    # offline (no live GigaChat call).
    def fake_structured(items, instruction, chunk_size=30):
        out = []
        for it in items:
            new = dict(it)
            if it["type"] == "text":
                new["text"] = it.get("text", "").replace("LXSHOW", "REMOVED")
            else:
                new["cells"] = [
                    [c.replace("LXSHOW", "REMOVED") if isinstance(c, str) else c for c in row]
                    for row in it.get("cells", [])
                ]
            out.append(new)
        return out

    monkeypatch.setattr(
        "app.services.document_edit._request_edits_structured", fake_structured
    )

    db = SessionLocal()
    try:
        result = edit_document(
            source_id, "Переведи на русский, убери LXSHOW, PDF", user_id, db
        )
    finally:
        db.close()

    assert result["success"] is True
    assert source_path.read_bytes() == content  # original untouched
    new_doc = (
        SessionLocal().query(Document).filter(Document.id == result["document_id"]).first()
    )
    out = fitz.open(stream=Path(new_doc.filepath).read_bytes(), filetype="pdf")
    src = fitz.open(stream=content, filetype="pdf")
    assert out.page_count == src.page_count == n_pages
    for pg in range(n_pages):
        assert len(out[pg].get_images()) == len(src[pg].get_images())
    assert "LXSHOW" not in out[0].get_text()
    out.close()
    src.close()


def test_pdf_translate_keeps_images_and_original(monkeypatch, user_id):
    content = _make_pdf_with_image("Original paragraph one.", "Original paragraph two.")
    source_id, source_path = _seed_pdf(user_id, content)

    # Deterministic 'translation' stub for the structured edit step: rewrite each
    # text item in place. Latin-only so the assertion is font-safe (base PDF
    # fonts lack Cyrillic).
    def fake_structured(items, instruction, chunk_size=30):
        out = []
        for it in items:
            new = dict(it)
            if it["type"] == "text":
                new["text"] = f"TRANSLATED {it.get('text', '')}"
            else:
                new["cells"] = [
                    [f"TRANSLATED {c}" if isinstance(c, str) else c for c in row]
                    for row in it.get("cells", [])
                ]
            out.append(new)
        return out

    monkeypatch.setattr(
        "app.services.document_edit._request_edits_structured", fake_structured
    )

    db = SessionLocal()
    try:
        result = edit_document(
            source_id, "Переведи на русский, убери LXSHOW, PDF", user_id, db
        )
    finally:
        db.close()

    assert result["success"] is True
    assert result["source_file_id"] == source_id
    assert result["file_type"] == "pdf"
    # Original on disk is byte-for-byte unchanged.
    assert source_path.read_bytes() == content

    new_doc = (
        SessionLocal().query(Document).filter(Document.id == result["document_id"]).first()
    )
    out = fitz.open(stream=Path(new_doc.filepath).read_bytes(), filetype="pdf")
    src = fitz.open(stream=content, filetype="pdf")
    # Page count and embedded images survive the edit.
    assert out.page_count == src.page_count
    assert len(out[0].get_images()) == len(src[0].get_images())
    assert "TRANSLATED" in out[0].get_text()
    out.close()
    src.close()


# ---------------------------------------------------------------- DOCX text-objects regression
# Reported bug: the parser rejected a model reply of {"text": ...} objects as
# soon as any block legitimately translated to an empty string, because the
# ``or``-chained lookup treated ``{"text": ""}`` as a missing field.


def test_request_edits_accepts_objects_with_empty_text(monkeypatch):
    """Exact reported shape: an array of {"text": ...} objects, some empty."""
    monkeypatch.setattr(
        gemini,
        "generate_answer",
        lambda prompt, system_instruction=None, **kwargs: (
            '[{"text":"ТЕХНИЧЕСКИЙ ОТЧЁТ"},{"text":"Клиент: ООО Альфа"},{"text":""}]'
        ),
    )
    assert _request_edits(["a", "b", "c"], "translate") == [
        "ТЕХНИЧЕСКИЙ ОТЧЁТ",
        "Клиент: ООО Альфа",
        "",
    ]


def test_parse_edits_array_accepts_object_with_empty_text():
    assert _parse_edits_array('[{"text":"one"},{"text":""}]', expected=2) == ["one", ""]
    assert _parse_edits_array('[{"text":""},{"text":"two"}]', expected=2) == ["", "two"]
    assert _parse_edits_array('[{"text":""},{"text":""}]', expected=2) == ["", ""]


def test_parse_edits_array_accepts_wrapped_object_array():
    """GigaChat's json_object mode wraps the array in a single object field."""
    assert (
        _parse_edits_array('{"edits": [{"text":"a"},{"text":"b"}]}', expected=2)
        == ["a", "b"]
    )


def test_parse_edits_array_accepts_mixed_strings_and_objects():
    assert _parse_edits_array('["a", {"text":"b"}, {"text":""}]', expected=3) == [
        "a",
        "b",
        "",
    ]


def test_parse_edits_array_rejects_object_count_mismatch():
    """Nine objects for eight blocks must be rejected, never silently trimmed."""
    assert (
        _parse_edits_array('[{"text":"a"},{"text":"b"},{"text":"c"}]', expected=2)
        is None
    )
    assert _parse_edits_array('[{"text":"a"}]', expected=2) is None


def test_parse_edits_array_rejects_object_missing_text():
    assert _parse_edits_array('[{"text":"a"},{"other":"b"}]', expected=2) is None


def test_request_edits_retries_on_wrong_count_and_accepts_objects(monkeypatch):
    """First attempt: 3 objects for two blocks (rejected). Retry: two valid
    objects including an empty one - success on the second call."""
    calls = {"n": 0}

    def fake(prompt, system_instruction=None, **kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            return '[{"text":"one"},{"text":"two"},{"text":"three"}]'
        return '[{"text":"one"},{"text":""}]'

    monkeypatch.setattr(gemini, "generate_answer", fake)
    assert _request_edits(["a", "b"], "translate") == ["one", ""]
    assert calls["n"] == 2


def test_request_edits_rejects_persistent_count_mismatch(monkeypatch):
    """An answer with 9 objects for 8 blocks is rejected on both attempts and
    surfaces as ``DocumentEditError`` - never silently accepted as 8."""
    payload = json.dumps([{"text": f"x{i}"} for i in range(9)])
    monkeypatch.setattr(
        gemini,
        "generate_answer",
        lambda prompt, system_instruction=None, **kwargs: payload,
    )
    with pytest.raises(DocumentEditError):
        _request_edits(["a"] * 8, "translate")


def _write_test_image() -> Path:
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 32, 32), False)
    png = pix.tobytes("png")
    path = upload_dir / f"img_{uuid.uuid4().hex}.png"
    path.write_bytes(png)
    return path


def _build_docx_bytes() -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("Технический отчёт", level=1)
    doc.add_paragraph("Исходный абзац один.")
    doc.add_paragraph("Клиент: ООО Альфа.")
    doc.add_paragraph("")  # empty block -> the model returns {"text": ""}
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Показатель"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Оклад"
    table.cell(1, 1).text = "150000"
    p_img = doc.add_paragraph()
    run = p_img.add_run("Подпись рисунка")
    run.add_picture(str(_write_test_image()))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _seed_docx(user_id: int, data: bytes) -> tuple[int, Path]:
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    stored = f"{uuid.uuid4().hex}.docx"
    path = upload_dir / stored
    path.write_bytes(data)
    db = SessionLocal()
    try:
        doc = Document(
            user_id=user_id,
            filename=stored,
            original_filename="manual.docx",
            file_type="docx",
            file_size=len(data),
            filepath=str(path),
            content="docx seed",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc.id, path
    finally:
        db.close()


def _docx_translate_stub(prompt, system_instruction=None, **kwargs):
    """Fake edit model mirroring the reported live behaviour: objects with a
    ``text`` field, where an empty source block becomes ``{"text": ""}``."""
    blob = (system_instruction or "") + "\n" + (prompt or "")
    texts = []
    for line in blob.splitlines():
        idx, sep, text = line.partition("\t")
        if sep and idx.isdigit():
            texts.append(text)
    out = []
    for t in texts:
        stripped = t.strip()
        out.append({"text": ("ПЕРЕВЕДЕНО: " + stripped) if stripped else ""})
    return json.dumps(out)


def test_docx_translate_real_scenario_preserves_images_tables_and_original(
    monkeypatch, user_id
):
    """Full DOCX translation via edit_document with the reported reply shape.

    Verifies the whole contract of the real scenario, not just the absence of
    an exception: text blocks translated, empty blocks accepted, tables and
    images preserved, the original byte-for-byte untouched, a separate result
    created and ``source_file_id`` pointing at the original.
    """
    from docx.oxml.ns import qn

    content = _build_docx_bytes()
    source_id, source_path = _seed_docx(user_id, content)

    monkeypatch.setattr(gemini, "generate_answer", _docx_translate_stub)

    db = SessionLocal()
    try:
        result = edit_document(
            source_id,
            "Переведи этот DOCX на русский язык, сохраняя форматирование, изображения и структуру документа.",
            user_id,
            db,
        )
    finally:
        db.close()

    assert result["success"] is True, result
    assert result["file_type"] == "docx"
    assert result["source_file_id"] == source_id
    # Original on disk is byte-for-byte unchanged.
    assert source_path.read_bytes() == content

    db = SessionLocal()
    try:
        new_doc = (
            db.query(Document)
            .filter(Document.id == result["document_id"])
            .first()
        )
    finally:
        db.close()
    assert new_doc is not None
    assert Path(new_doc.filepath) != source_path  # a separate result file

    from docx import Document as DocxDocument

    out = DocxDocument(Path(new_doc.filepath))
    paras = [p.text for p in out.paragraphs]
    assert "ПЕРЕВЕДЕНО: Исходный абзац один." in paras
    assert "ПЕРЕВЕДЕНО: Технический отчёт" in paras or any(
        "ПЕРЕВЕДЕНО: Технический отчёт" in t for t in paras
    )
    assert "ПЕРЕВЕДЕНО: Подпись рисунка" in paras

    # Structure preserved: the table and the embedded image survive.
    assert len(out.tables) == 1
    tbl = out.tables[0]
    assert len(tbl.rows) == 2 and len(tbl.columns) == 2
    assert tbl.cell(0, 0).text.startswith("ПЕРЕВЕДЕНО:")
    drawings = list(out.element.body.iter(qn("w:drawing")))
    assert len(drawings) >= 1
